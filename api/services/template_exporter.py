"""
Template Exporter - Export templates to HTML, Markdown, DOCX formats

Supports platform-specific export formats:
- Saramin (사람인): Template 2 (경력) + Template 3 (경력기술서)
- Remember (리멤버): Template 4 (회사+프로젝트 통합)
- Wanted (원티드): Template 3 (프로젝트 상세만)
- Jumpit (점핏): Template 3 (프로젝트 상세만)
"""

import os
import re
from typing import Tuple, Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

from api.schemas.platform import RenderDataRequest
from api.config import get_settings
from api.services.docx_styles import DocxStyler, DocxStyleConfig, DEFAULT_DOCX_STYLE

try:
    import chevron
except ImportError:
    chevron = None

settings = get_settings()

# Domain categorization keywords for Template 2 style
DOMAIN_CATEGORIES = {
    "Backend": {
        "keywords": ["python", "fastapi", "django", "flask", "spring", "java", "kotlin", "express", "nodejs", "restful", "api", "jwt", "oauth"],
        "name_ko": "Backend",
    },
    "AI/ML": {
        "keywords": ["tensorflow", "pytorch", "scikit", "langchain", "llm", "gpt", "openai", "ml", "ai", "machine learning", "deep learning", "rag", "nlp"],
        "name_ko": "AI/ML",
    },
    "Frontend": {
        "keywords": ["react", "vue", "angular", "next", "typescript", "javascript", "html", "css", "tailwind", "shadcn"],
        "name_ko": "Frontend",
    },
    "Mobile": {
        "keywords": ["flutter", "dart", "swift", "kotlin", "react native", "android", "ios", "mobile"],
        "name_ko": "Mobile/Cross-Platform",
    },
    "Data": {
        "keywords": ["pandas", "numpy", "matplotlib", "data analysis", "데이터 분석", "excel", "visualization"],
        "name_ko": "데이터 분석",
    },
    "Database": {
        "keywords": ["postgresql", "mysql", "sqlite", "mongodb", "redis", "elasticsearch", "sql", "database", "db"],
        "name_ko": "Database",
    },
    "DevOps": {
        "keywords": ["docker", "kubernetes", "aws", "gcp", "azure", "ci/cd", "github action", "nginx", "linux"],
        "name_ko": "DevOps/인프라",
    },
    "IoT": {
        "keywords": ["iot", "embedded", "ble", "bluetooth", "sensor", "raspberry", "arduino", "임베디드"],
        "name_ko": "IoT/임베디드",
    },
}


class TemplateExporter:
    """Handles exporting templates to various formats"""

    def __init__(
        self,
        result_dir: Path = None,
        style_config: Optional[DocxStyleConfig] = None
    ):
        self.result_dir = result_dir or Path(settings.result_dir)
        self.styler = DocxStyler(style_config or DEFAULT_DOCX_STYLE)

    def generate_filename(self, platform_key: str, name: str, ext: str) -> str:
        """Generate a safe filename with timestamp"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_name = re.sub(r'[^\w\s-]', '', name or "resume").strip().replace(' ', '_')[:20]
        return f"{platform_key}_{safe_name}_{timestamp}.{ext}"

    def save_file(self, filename: str, content: str) -> str:
        """Save content to file and return path"""
        os.makedirs(self.result_dir, exist_ok=True)
        file_path = self.result_dir / filename
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return str(file_path)

    def export_html(
        self,
        html_content: str,
        platform_key: str,
        name: str
    ) -> Tuple[str, str]:
        """Export to HTML file"""
        filename = self.generate_filename(platform_key, name, "html")
        file_path = self.save_file(filename, html_content)
        return file_path, html_content

    def export_markdown(
        self,
        data: RenderDataRequest,
        platform_key: str,
        template_name: str
    ) -> Tuple[str, str]:
        """Export to Markdown file"""
        md_content = self.generate_markdown(data, template_name)
        filename = self.generate_filename(platform_key, data.name, "md")
        file_path = self.save_file(filename, md_content)
        return file_path, md_content

    def export_docx(
        self,
        data: RenderDataRequest,
        platform_key: str
    ) -> str:
        """Export to Word document with proper styling

        Uses DocxStyler for consistent formatting:
        - Title (대제목): 24pt, bold, black
        - Heading1 (중제목): 18pt, bold, black
        - Heading2 (소제목): 14pt, bold, black
        - Normal text: 11pt, black
        """
        from docx import Document

        filename = self.generate_filename(platform_key, data.name, "docx")
        os.makedirs(self.result_dir, exist_ok=True)
        file_path = self.result_dir / filename

        doc = Document()

        # Set up document styles (removes colors, sets fonts)
        self.styler.setup_document(doc)

        # Header - Title (대제목)
        self.styler.add_title(doc, data.name or "이력서")

        if data.desired_position:
            self.styler.add_subtitle(doc, data.desired_position)

        # Contact info
        contact_parts = []
        if data.email:
            contact_parts.append(data.email)
        if data.phone:
            contact_parts.append(data.phone)
        if data.github_url:
            contact_parts.append(data.github_url)

        self.styler.add_contact_info(doc, contact_parts)
        self.styler.add_spacing(doc)

        # Summary
        if data.summary:
            self.styler.add_heading1(doc, "자기소개")
            self.styler.add_paragraph(doc, data.summary)

        # Experience
        if data.experiences:
            self.styler.add_heading1(doc, "경력사항")
            for exp in data.experiences:
                # Company name as heading2 (소제목)
                period = f" ({exp.start_date or ''} ~ {exp.end_date or '현재'})"
                self.styler.add_heading2(doc, f"{exp.company_name}{period}")

                # Build context for Mustache rendering
                exp_context = {
                    "company_name": exp.company_name,
                    "position": exp.position,
                    "start_date": exp.start_date,
                    "end_date": exp.end_date,
                    "description": exp.description,
                    "achievements": exp.achievements,
                }

                if exp.position:
                    self.styler.add_bullet(doc, exp.position)
                if exp.description:
                    # Clean any Mustache syntax from description
                    clean_desc = self._render_or_clean_mustache(exp.description, exp_context)
                    if clean_desc:
                        self.styler.add_paragraph(doc, clean_desc)
                if exp.achievements:
                    for ach in exp.achievements:
                        # Clean any Mustache syntax from achievements
                        clean_ach = self._render_or_clean_mustache(ach, exp_context)
                        if clean_ach:
                            self.styler.add_bullet(doc, clean_ach)

        # Projects
        if data.projects:
            self.styler.add_heading1(doc, "프로젝트")
            for proj in data.projects:
                # Project name as heading2 (소제목)
                period = f" ({proj.start_date or ''} ~ {proj.end_date or '현재'})"
                self.styler.add_heading2(doc, f"{proj.name}{period}")

                # Build context for Mustache rendering
                proj_context = {
                    "name": proj.name,
                    "company_name": proj.company_name,
                    "start_date": proj.start_date,
                    "end_date": proj.end_date,
                    "description": proj.description,
                    "role": proj.role,
                    "technologies": proj.technologies,
                    "achievements": proj.achievements,
                }

                if proj.company_name:
                    self.styler.add_paragraph(doc, f"소속: {proj.company_name}")
                if proj.description:
                    # Clean any Mustache syntax from description
                    clean_desc = self._render_or_clean_mustache(proj.description, proj_context)
                    if clean_desc:
                        self.styler.add_paragraph(doc, clean_desc)
                if proj.role:
                    # Clean any Mustache syntax from role
                    clean_role = self._render_or_clean_mustache(proj.role, proj_context)
                    if clean_role:
                        self.styler.add_paragraph(doc, f"역할: {clean_role}")
                if proj.technologies:
                    self.styler.add_paragraph(doc, f"기술: {', '.join(proj.technologies)}")
                # Use achievements_summary_list (from detailed_achievements) as default, fall back to achievements
                achievements_list = self._get_achievements_list({
                    "achievements_summary_list": proj.achievements_summary_list,
                    "achievements_detailed_list": proj.achievements_detailed_list,
                    "achievements": proj.achievements,
                })
                if achievements_list:
                    for ach in achievements_list:
                        # Clean any Mustache syntax from achievements
                        clean_ach = self._render_or_clean_mustache(ach, proj_context)
                        if clean_ach:
                            self.styler.add_bullet(doc, clean_ach)

        # Skills
        if data.skills:
            self.styler.add_heading1(doc, "기술 스택")
            skills = data.skills
            if skills.languages:
                self.styler.add_key_value(doc, "Languages", ", ".join(skills.languages))
            if skills.frameworks:
                self.styler.add_key_value(doc, "Frameworks", ", ".join(skills.frameworks))
            if skills.databases:
                self.styler.add_key_value(doc, "Databases", ", ".join(skills.databases))
            if skills.tools:
                self.styler.add_key_value(doc, "Tools", ", ".join(skills.tools))

        # Education
        if data.educations:
            self.styler.add_heading1(doc, "학력")
            for edu in data.educations:
                period = f" ({edu.start_date or ''} ~ {edu.end_date or ''})"
                major_str = f" - {edu.major}" if edu.major else ""
                self.styler.add_heading2(doc, f"{edu.school_name}{major_str}{period}")

        # Certifications
        if data.certifications:
            self.styler.add_heading1(doc, "자격증")
            for cert in data.certifications:
                self.styler.add_bullet(doc, f"{cert.name} ({cert.issuer or ''}) - {cert.date or ''}")

        doc.save(file_path)
        return str(file_path)

    def generate_markdown(
        self,
        data: RenderDataRequest,
        template_name: str,
        platform_key: str = None
    ) -> str:
        """Generate Markdown from RenderDataRequest user data

        Args:
            data: RenderDataRequest Pydantic model
            template_name: Template name for footer
            platform_key: Platform key for format selection

        Note: For platform-specific formatting with dict data, use generate_markdown_from_dict
        """
        # Convert RenderDataRequest to dict for platform-specific generation
        data_dict = self._render_data_to_dict(data)
        return self.generate_markdown_from_dict(data_dict, template_name, platform_key)

    def _render_data_to_dict(self, data: RenderDataRequest) -> Dict[str, Any]:
        """Convert RenderDataRequest Pydantic model to dict for processing"""
        result = {
            "name": data.name,
            "email": data.email,
            "phone": data.phone,
            "photo_url": data.photo_url,
            "desired_position": data.desired_position,
            "summary": data.summary,
            "introduction": data.summary,  # Alias
            "github_url": data.github_url,
            "portfolio_url": data.portfolio_url,
        }

        # Convert experiences
        if data.experiences:
            result["experiences"] = [
                {
                    "company_name": exp.company_name,
                    "position": exp.position,
                    "start_date": exp.start_date,
                    "end_date": exp.end_date,
                    "description": exp.description,
                    "achievements": exp.achievements,
                }
                for exp in data.experiences
            ]
        else:
            result["experiences"] = []

        # Convert projects
        if data.projects:
            result["projects"] = [
                {
                    "name": proj.name,
                    "company_name": proj.company_name,
                    "company": proj.company_name,  # Alias
                    "start_date": proj.start_date,
                    "end_date": proj.end_date,
                    "description": proj.description,
                    "role": proj.role,
                    "technologies": proj.technologies,
                    "technologies_list": proj.technologies,
                    "achievements": proj.achievements,
                    # Extended fields for platform-specific exports
                    "key_tasks_list": proj.key_tasks_list or [],
                    "team_size": proj.team_size,
                    "implementation_details": proj.implementation_details or [],
                    "achievements_summary_list": proj.achievements_summary_list or [],  # [{category, title}]
                    "achievements_detailed_list": proj.achievements_detailed_list or [],  # [{category, title, description}]
                }
                for proj in data.projects
            ]
        else:
            result["projects"] = []

        # Convert skills
        if data.skills:
            result["skills_categorized"] = {
                "languages": data.skills.languages or [],
                "frameworks": data.skills.frameworks or [],
                "databases": data.skills.databases or [],
                "tools": data.skills.tools or [],
            }
            # Flatten skills list
            result["skills"] = (
                (data.skills.languages or []) +
                (data.skills.frameworks or []) +
                (data.skills.databases or []) +
                (data.skills.tools or [])
            )
        else:
            result["skills_categorized"] = {}
            result["skills"] = []

        # Convert education
        if data.educations:
            result["education"] = [
                {
                    "school_name": edu.school_name,
                    "major": edu.major,
                    "degree": edu.degree,
                    "start_date": edu.start_date,
                    "end_date": edu.end_date,
                    "description": edu.description,
                }
                for edu in data.educations
            ]
        else:
            result["education"] = []

        # Convert certifications
        if data.certifications:
            result["certifications"] = [
                {
                    "name": cert.name,
                    "issuer": cert.issuer,
                    "date": cert.date,
                }
                for cert in data.certifications
            ]
        else:
            result["certifications"] = []

        return result

    def export_markdown_from_dict(
        self,
        data: dict,
        platform_key: str,
        template_name: str
    ) -> Tuple[str, str]:
        """Export to Markdown file from dict data with platform-specific format"""
        md_content = self.generate_markdown_from_dict(data, template_name, platform_key)
        filename = self.generate_filename(platform_key, data.get("name", "user"), "md")
        file_path = self.save_file(filename, md_content)
        return file_path, md_content

    def export_docx_from_dict(
        self,
        data: dict,
        platform_key: str
    ) -> str:
        """Export to Word document from dict data with unified format

        All platforms use the same Word format:
        - 경력: Company-based with domain-categorized skills
        - 경력기술서: Detailed project information
        """
        return self._export_unified_docx(data, platform_key)

    def _export_unified_docx(self, data: dict, platform_key: str) -> str:
        """Unified Word format for all platforms with proper styling

        Structure:
        - 경력: Company-based with domain-categorized skills
        - 경력기술서: Detailed project information

        Styling:
        - Title (대제목): 24pt, bold, black
        - Heading1 (중제목): 18pt, bold, black
        - Heading2 (소제목): 14pt, bold, black
        - Normal text: 11pt, black
        """
        from docx import Document

        filename = self.generate_filename(platform_key, data.get("name", "user"), "docx")
        os.makedirs(self.result_dir, exist_ok=True)
        file_path = self.result_dir / filename

        doc = Document()

        # Set up document styles (removes colors, sets fonts)
        self.styler.setup_document(doc)

        # Header - Title (대제목)
        self.styler.add_title(doc, data.get("name", "이력서"))

        # Contact info
        self._add_contact_info_docx(doc, data)

        experiences = data.get("experiences", [])
        projects = data.get("projects", [])

        # Calculate total experience
        total_years = data.get("total_experience_years", 0)

        # === Part 1: 경력 ===
        if experiences or projects:
            heading_text = f"경력 (총 {total_years}년)" if total_years else "경력"
            self.styler.add_heading1(doc, heading_text)

            company_projects = self._group_projects_by_company(projects)

            for exp in experiences:
                company_name = exp.get("company_name", "")
                position = exp.get("position", "")
                start_date = exp.get("start_date", "")
                end_date = exp.get("end_date", "")
                is_current = exp.get("is_current", not end_date)
                duration = exp.get("duration", "")
                description = exp.get("description", "")

                # Company header (소제목)
                self.styler.add_heading2(doc, company_name)

                if position:
                    self.styler.add_paragraph(doc, position)

                period_str = f"{start_date} ~ {end_date if end_date else ''}"
                if is_current:
                    period_str += " 재직중" if not end_date else ""
                if duration:
                    period_str += f" ({duration})"
                self.styler.add_paragraph(doc, period_str)

                if description:
                    self.styler.add_bullet(doc, f"주요 직무: {description}")

                # Domain skills
                company_projs = company_projects.get(company_name, [])
                skills_by_domain = self._categorize_skills_by_domain_detailed(company_projs)

                domain_idx = 1
                for domain_name, domain_data in skills_by_domain.items():
                    if not domain_data.get("technologies") and not domain_data.get("implementations"):
                        continue

                    self.styler.add_bold_text(doc, f"{domain_idx}. {domain_name}")

                    if domain_data.get("technologies"):
                        self.styler.add_paragraph(doc, "주요 사용 기술")
                        for tech_line in domain_data["technologies"]:
                            self.styler.add_bullet(doc, tech_line)

                    if domain_data.get("implementations"):
                        self.styler.add_paragraph(doc, "구현 내용")
                        for impl in domain_data["implementations"]:
                            self.styler.add_bullet(doc, impl)

                    if domain_data.get("databases"):
                        self.styler.add_paragraph(doc, "DB")
                        for db_line in domain_data["databases"]:
                            self.styler.add_bullet(doc, db_line)

                    domain_idx += 1

                self.styler.add_spacing(doc)

            # Side projects
            side_projects = company_projects.get(None, []) + company_projects.get("", []) + company_projects.get("사이드 프로젝트", [])
            if side_projects:
                self.styler.add_heading2(doc, "사이드 프로젝트")

                skills_by_domain = self._categorize_skills_by_domain_detailed(side_projects)

                domain_idx = 1
                for domain_name, domain_data in skills_by_domain.items():
                    if not domain_data.get("implementations"):
                        continue

                    self.styler.add_bold_text(doc, f"{domain_idx}. {domain_name}")

                    for impl in domain_data["implementations"]:
                        self.styler.add_bullet(doc, impl)

                    domain_idx += 1

        # === Part 2: 경력기술서 ===
        if projects:
            self.styler.add_heading1(doc, "경력기술서")

            for idx, proj in enumerate(projects, 1):
                self._add_project_detail_unified_docx(doc, proj, idx)

        doc.save(file_path)
        return str(file_path)

    def _add_contact_info_docx(self, doc, data: dict) -> None:
        """Add contact info to Word document using styler"""
        contact_parts = []
        if data.get("email"):
            contact_parts.append(data["email"])
        if data.get("phone"):
            contact_parts.append(data["phone"])
        if data.get("github_url"):
            contact_parts.append(data["github_url"])

        self.styler.add_contact_info(doc, contact_parts)
        self.styler.add_spacing(doc)

    def _add_project_detail_unified_docx(self, doc, proj: dict, idx: int) -> None:
        """Add project in unified format to Word document using styler

        Note: Uses _render_or_clean_mustache to handle any Mustache template
        syntax that may be in the text fields (e.g., {{key_tasks}}, {{#achievements}})
        """
        # Project header (소소제목)
        self.styler.add_section_title(doc, f"[프로젝트 {idx}]")

        self.styler.add_key_value(doc, "프로젝트명", proj.get('name', ''))

        company = proj.get("company_name") or proj.get("company")
        if company:
            self.styler.add_key_value(doc, "연계/소속회사", company)
        else:
            self.styler.add_key_value(doc, "연계/소속회사", "사이드 프로젝트")

        end_date = proj.get("end_date") or "진행중"
        self.styler.add_key_value(doc, "기간", f"{proj.get('start_date', '')} ~ {end_date}")

        role = proj.get("role")
        if role:
            # Clean any Mustache syntax from role
            clean_role = self._render_or_clean_mustache(role, proj)
            if clean_role:
                self.styler.add_key_value(doc, "주요 업무", clean_role)

        # 주요 구현 기능 (최대 8개)
        key_tasks = self._get_key_tasks_list(proj)
        if key_tasks:
            self.styler.add_paragraph(doc, "주요 구현 기능:")
            for task in key_tasks[:8]:
                # Clean any Mustache syntax from each task
                clean_task = self._render_or_clean_mustache(task, proj)
                if clean_task:
                    self.styler.add_bullet(doc, clean_task)

        # 기술 스택
        tech = proj.get("technologies")
        if tech:
            tech_str = tech if isinstance(tech, str) else ", ".join(tech)
            self.styler.add_key_value(doc, "기술 스택", tech_str)

        # 개발 인원
        team_size = proj.get("team_size")
        if team_size:
            self.styler.add_key_value(doc, "개발 인원", str(team_size))

        # 상세 내용 - clean any Mustache template syntax
        description = proj.get("description", "")
        if description:
            # Render or clean Mustache syntax from description
            clean_description = self._render_or_clean_mustache(description, proj)
            if clean_description:
                self.styler.add_key_value(doc, "상세 내용", clean_description)

        # 성과 (최대 4개)
        achievements = self._get_achievements_list(proj)
        if achievements:
            self.styler.add_paragraph(doc, "성과:")
            for ach in achievements[:4]:
                # Clean any Mustache syntax from each achievement
                clean_ach = self._render_or_clean_mustache(ach, proj)
                if clean_ach:
                    self.styler.add_bullet(doc, clean_ach)

        self.styler.add_spacing(doc)

    def generate_markdown_from_dict(
        self,
        data: dict,
        template_name: str,
        platform_key: str = None
    ) -> str:
        """Generate unified Markdown from dict data

        All platforms use the same Markdown/Word format:
        - 경력: Company-based with domain-categorized skills
        - 경력기술서: Detailed project information

        Args:
            data: User data dictionary
            template_name: Template name for footer
            platform_key: Platform key (not used for format selection, only for footer)

        Returns:
            Markdown string with unified format
        """
        # All platforms use the same unified format
        return self._generate_unified_markdown(data, template_name)

    def _generate_unified_markdown(self, data: dict, template_name: str) -> str:
        """Generate unified Markdown format for all platforms

        Structure (based on user's detailed example):
        1. 경력: Company-based with domain-categorized skills
        2. 경력기술서: Detailed project information

        Format follows user's exact specification with:
        - Company header with period, position, job type
        - Domain sections (Backend, AI/ML, Mobile, etc.) with:
          - 주요 사용 기술
          - 구현 내용
          - DB (if applicable)
        - Project details with bullet points
        """
        lines = []

        # === Header ===
        lines.append(f"# {data.get('name', '이력서')}")
        lines.append("")

        # Contact info
        self._append_contact_info(lines, data)

        lines.append("---")
        lines.append("")

        experiences = data.get("experiences", [])
        projects = data.get("projects", [])

        # Calculate total experience
        total_years = data.get("total_experience_years", 0)
        if total_years:
            lines.append(f"## 경력")
            lines.append(f"총 {total_years}년")
            lines.append("")
        elif experiences:
            lines.append("## 경력")
            lines.append("")

        # Group projects by company
        company_projects = self._group_projects_by_company(projects)

        # === Part 1: 경력 (Company-based with domain skills) ===
        if experiences:
            for exp in experiences:
                company_name = exp.get("company_name", "")
                position = exp.get("position", "")
                start_date = exp.get("start_date", "")
                end_date = exp.get("end_date", "")
                is_current = exp.get("is_current", not end_date)
                duration = exp.get("duration", "")
                description = exp.get("description", "")

                # Company header
                lines.append(f"### {company_name}")
                if position:
                    lines.append(f"{position}")
                period_str = f"{start_date} ~ {end_date if end_date else ''}"
                if is_current:
                    period_str += " 재직중" if not end_date else ""
                if duration:
                    period_str += f" ({duration})"
                lines.append(period_str)
                lines.append("")

                # Main duties
                if description:
                    lines.append(f"- 주요 직무: {description}")
                    lines.append("")

                # Categorize skills by domain for this company's projects
                company_projs = company_projects.get(company_name, [])
                skills_by_domain = self._categorize_skills_by_domain_detailed(company_projs)

                domain_idx = 1
                for domain_name, domain_data in skills_by_domain.items():
                    if not domain_data.get("technologies") and not domain_data.get("implementations"):
                        continue

                    lines.append(f"**{domain_idx}. {domain_name}**")

                    # Technologies
                    if domain_data.get("technologies"):
                        lines.append("주요 사용 기술")
                        for tech_line in domain_data["technologies"]:
                            lines.append(f"- {tech_line}")

                    # Implementations
                    if domain_data.get("implementations"):
                        lines.append("")
                        lines.append("구현 내용")
                        for impl in domain_data["implementations"]:
                            lines.append(f"- {impl}")

                    # Databases
                    if domain_data.get("databases"):
                        lines.append("")
                        lines.append("DB")
                        for db_line in domain_data["databases"]:
                            lines.append(f"- {db_line}")

                    lines.append("")
                    domain_idx += 1

                lines.append("")

        # Handle side projects (projects without company in experiences)
        side_projects = company_projects.get(None, []) + company_projects.get("", []) + company_projects.get("사이드 프로젝트", [])
        if side_projects:
            lines.append("### 사이드 프로젝트")
            lines.append("")

            skills_by_domain = self._categorize_skills_by_domain_detailed(side_projects)

            domain_idx = 1
            for domain_name, domain_data in skills_by_domain.items():
                if not domain_data.get("technologies") and not domain_data.get("implementations"):
                    continue

                lines.append(f"**{domain_idx}. {domain_name}**")

                if domain_data.get("implementations"):
                    for impl in domain_data["implementations"]:
                        lines.append(f"- {impl}")

                lines.append("")
                domain_idx += 1

            lines.append("")

        # === Part 2: 경력기술서 (Project Details) ===
        if projects:
            lines.append("---")
            lines.append("")
            lines.append("## 경력기술서")
            lines.append("")

            for idx, proj in enumerate(projects, 1):
                self._append_project_detail_unified(lines, proj, idx)

        # Footer
        self._append_footer(lines, template_name)

        return "\n".join(lines)

    # ==================== Helper Methods ====================

    def _render_or_clean_mustache(self, text: str, context: dict = None) -> str:
        """Render Mustache template syntax or clean it from text

        Handles cases like:
        - {{key_tasks}} -> renders with context or removes
        - {{#achievements}}...{{/achievements}} -> renders loop or removes section

        Args:
            text: Text that may contain Mustache syntax
            context: Data context for rendering (e.g., project dict)

        Returns:
            Cleaned text with Mustache syntax rendered or removed
        """
        if not text:
            return ""

        # Check if text contains Mustache syntax
        if "{{" not in text:
            return text

        # Try to render with chevron if available and context provided
        if chevron and context:
            try:
                rendered = chevron.render(text, context)
                # Clean up any remaining empty lines from unresolved sections
                lines = [line for line in rendered.split('\n') if line.strip()]
                return '\n'.join(lines)
            except Exception:
                pass

        # Fallback: Remove Mustache syntax patterns
        import re

        # Remove section blocks: {{#section}}...{{/section}} (including multiline)
        result = re.sub(r'\{\{[#^]\w+\}\}.*?\{\{/\w+\}\}', '', text, flags=re.DOTALL)

        # Remove unclosed/malformed section blocks: {{#section}}... or {{/section}}
        result = re.sub(r'\{\{[#^/]\w+\}\}', '', result)

        # Remove single variables: {{variable}}, {{&variable}}, {{{variable}}}
        result = re.sub(r'\{{2,3}[&]?\w+\}{2,3}', '', result)

        # Remove partial includes: {{>partial}}
        result = re.sub(r'\{\{>\w+\}\}', '', result)

        # Remove any remaining {{ or }} fragments
        result = re.sub(r'\{\{[^}]*\}?\}?', '', result)

        # Clean up lines that only have bullet points or headers without content
        lines = []
        for line in result.split('\n'):
            stripped = line.strip()
            # Skip empty lines and lines with only bullet point markers
            if stripped and stripped not in ['•', '-', '*', '####', '###', '##', '#']:
                # Skip lines that are just headers with no following content
                if stripped.startswith('#') and ':' not in stripped and len(stripped) < 50:
                    continue
                lines.append(line)

        # Clean up multiple consecutive empty lines
        result = '\n'.join(lines)
        result = re.sub(r'\n\s*\n\s*\n', '\n\n', result)

        return result.strip()

    def _append_contact_info(self, lines: list, data: dict) -> None:
        """Append contact information to lines"""
        contact = []
        if data.get("email"):
            contact.append(data['email'])
        if data.get("phone"):
            contact.append(data['phone'])
        if data.get("github_url"):
            contact.append(data['github_url'])
        if data.get("portfolio_url") and data.get("portfolio_url") != data.get("github_url"):
            contact.append(data['portfolio_url'])

        if contact:
            lines.append(" | ".join(contact))
            lines.append("")

    def _append_footer(self, lines: list, template_name: str) -> None:
        """Append footer to lines"""
        lines.append("---")
        lines.append("")
        lines.append(f"*Generated by Autopolio ({template_name}) | {datetime.now().strftime('%Y-%m-%d')}*")

    def _append_project_detail_unified(self, lines: list, proj: dict, idx: int) -> None:
        """Append project in unified format (based on user's exact example)

        Format:
        [프로젝트 N]
        프로젝트명: ...
        연계/소속회사: ...
        기간: ...
        주요 업무: ...
        주요 구현 기능:
        • ...
        기술 스택: ...
        개발 인원: ...
        상세 내용: ...
        성과:
        • ...

        Note: Uses _render_or_clean_mustache to handle any Mustache template
        syntax that may be in the text fields (e.g., {{key_tasks}}, {{#achievements}})
        """
        lines.append(f"[프로젝트 {idx}]")
        lines.append(f"프로젝트명: {proj.get('name', '')}")

        company = proj.get("company_name") or proj.get("company")
        if company:
            lines.append(f"연계/소속회사: {company}")
        else:
            lines.append("연계/소속회사: 사이드 프로젝트")

        end_date = proj.get("end_date") or "진행중"
        lines.append(f"기간: {proj.get('start_date', '')} ~ {end_date}")

        role = proj.get("role")
        if role:
            # Clean any Mustache syntax from role
            clean_role = self._render_or_clean_mustache(role, proj)
            if clean_role:
                lines.append(f"주요 업무: {clean_role}")

        # 주요 구현 기능 (최대 8개)
        key_tasks = self._get_key_tasks_list(proj)
        if key_tasks:
            lines.append("주요 구현 기능:")
            for task in key_tasks[:8]:
                # Clean any Mustache syntax from each task
                clean_task = self._render_or_clean_mustache(task, proj)
                if clean_task:
                    lines.append(f"• {clean_task}")

        # 기술 스택
        tech = proj.get("technologies")
        if tech:
            tech_str = tech if isinstance(tech, str) else ", ".join(tech)
            lines.append(f"기술 스택: {tech_str}")

        # 개발 인원
        team_size = proj.get("team_size")
        if team_size:
            lines.append(f"개발 인원: {team_size}")

        # 상세 내용 - clean any Mustache template syntax
        description = proj.get("description", "")
        if description:
            # Render or clean Mustache syntax from description
            clean_description = self._render_or_clean_mustache(description, proj)
            if clean_description:
                lines.append(f"상세 내용: {clean_description}")

        # 성과 (최대 4개)
        achievements = self._get_achievements_list(proj)
        if achievements:
            lines.append("성과:")
            for ach in achievements[:4]:
                # Clean any Mustache syntax from each achievement
                clean_ach = self._render_or_clean_mustache(ach, proj)
                if clean_ach:
                    lines.append(f"• {clean_ach}")

        lines.append("")

    def _get_key_tasks_list(self, proj: dict) -> List[str]:
        """Extract key tasks as a list of strings"""
        key_tasks = proj.get("key_tasks_list", [])
        if not key_tasks:
            key_tasks_str = proj.get("key_tasks", "")
            if key_tasks_str:
                key_tasks = [t.strip().lstrip("•").strip()
                            for t in key_tasks_str.split("\n") if t.strip()]
        return key_tasks

    def _get_achievements_list(self, proj: dict, use_detailed: bool = False) -> List[str]:
        """Extract achievements as a list of strings

        Args:
            proj: Project dictionary
            use_detailed: If True, use achievements_detailed_list (with descriptions)
                         If False (default), use achievements_summary_list (title only)

        Priority order (DEFAULT = Summary):
        1. achievements_summary_list (from RepoAnalysis.detailed_achievements) - title only, no description
        2. achievements_basic_list (from ProjectAchievement) - metric_name: metric_value

        Returns formatted strings like:
        - Summary: "[새로운 기능 추가] UI 개선"
        - Detailed: "[새로운 기능 추가] UI 개선 - 접이식 문서 메뉴 추가, 사용자 경험 개선"
        """
        # If detailed is requested, use achievements_detailed_list first
        if use_detailed:
            detailed_list = proj.get("achievements_detailed_list", [])
            if detailed_list and isinstance(detailed_list, list):
                result = []
                for ach in detailed_list:
                    if isinstance(ach, dict):
                        category = ach.get("category", "")
                        title = ach.get("title", "")
                        description = ach.get("description", "")

                        if title:
                            if category:
                                line = f"[{category}] {title}"
                            else:
                                line = title
                            # Add description if available
                            if description:
                                line += f" - {description}"
                            result.append(line)
                if result:
                    return result

        # Priority 1: Use achievements_summary_list (from detailed_achievements) - DEFAULT
        # Format: [{category, title}] - NO description
        summary_list = proj.get("achievements_summary_list", [])
        if summary_list and isinstance(summary_list, list):
            result = []
            for ach in summary_list:
                if isinstance(ach, dict):
                    category = ach.get("category", "")
                    title = ach.get("title", "")

                    if title:
                        if category:
                            line = f"[{category}] {title}"
                        else:
                            line = title
                        result.append(line)
            if result:
                return result

        # Priority 2: Use achievements_basic_list (from ProjectAchievement)
        basic_list = proj.get("achievements_basic_list", [])
        if basic_list and isinstance(basic_list, list):
            result = []
            for ach in basic_list:
                if isinstance(ach, dict):
                    metric = ach.get("metric_name", "")
                    value = ach.get("metric_value", "")
                    if metric and value:
                        result.append(f"{metric}: {value}")
                    elif metric:
                        result.append(metric)
            if result:
                return result

        # Fallback: Use achievements string
        achievements = proj.get("achievements")
        if isinstance(achievements, str):
            return [a.strip().lstrip("•").lstrip("-").strip()
                    for a in achievements.split("\n") if a.strip()]
        return []

    def _group_projects_by_company(self, projects: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """Group projects by company name"""
        company_projects = {}
        for proj in projects:
            company = proj.get("company_name") or proj.get("company") or ""
            if company not in company_projects:
                company_projects[company] = []
            company_projects[company].append(proj)
        return company_projects

    def _categorize_skills_by_domain_detailed(self, projects: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
        """Categorize skills by domain with detailed format (based on user's example)

        Returns dict with:
        - technologies: List of technology lines (can be multiple items per line)
        - implementations: List of implementation details
        - databases: List of database-related items

        Example output for Backend domain:
        {
            "technologies": [
                "FastAPI, Spring Boot, RESTful API, JWT 인증",
                "ORM (SQLAlchemy, JPA/Hibernate)",
                "Linux (Rocky Linux, Ubuntu), Nginx, Docker"
            ],
            "implementations": [
                "RESTful API 서버 개발 (10개 이상 프로젝트)",
                "데이터 수집 파이프라인 구축",
                ...
            ],
            "databases": [
                "MySQL, PostgreSQL, SQLite, Elasticsearch",
                "DB 설계, ERD 작성, 테이블 정의서 작성"
            ]
        }
        """
        domain_data = {}

        for proj in projects:
            # Get technologies
            tech = proj.get("technologies")
            if isinstance(tech, str):
                techs = [t.strip() for t in tech.split(",") if t.strip()]
            elif isinstance(tech, list):
                techs = tech
            else:
                techs = []

            # Get key tasks and implementation details
            key_tasks = self._get_key_tasks_list(proj)
            implementation_details = proj.get("implementation_details", [])

            # Find primary domain
            project_primary_domain = None
            domain_tech_count = {}

            # Categorize technologies
            for tech_name in techs:
                tech_lower = tech_name.lower()
                matched_domain = None

                for domain_key, domain_info in DOMAIN_CATEGORIES.items():
                    if any(kw in tech_lower for kw in domain_info["keywords"]):
                        matched_domain = domain_info["name_ko"]
                        break

                if not matched_domain:
                    matched_domain = "기타"

                if matched_domain not in domain_data:
                    domain_data[matched_domain] = {
                        "technologies": [],
                        "implementations": [],
                        "databases": [],
                        "_tech_set": set(),
                        "_impl_set": set(),
                        "_db_set": set(),
                    }

                domain_data[matched_domain]["_tech_set"].add(tech_name)
                domain_tech_count[matched_domain] = domain_tech_count.get(matched_domain, 0) + 1

                # Check if it's a database tech
                if any(kw in tech_lower for kw in DOMAIN_CATEGORIES.get("Database", {}).get("keywords", [])):
                    domain_data[matched_domain]["_db_set"].add(tech_name)

            # Find primary domain
            if domain_tech_count:
                project_primary_domain = max(domain_tech_count.items(), key=lambda x: x[1])[0]

            # Process implementation_details from LLM analysis
            for detail in implementation_details:
                if isinstance(detail, dict):
                    title = detail.get("title", "")
                    items = detail.get("items", [])
                    title_lower = title.lower()

                    matched_domain = None
                    for domain_key, domain_info in DOMAIN_CATEGORIES.items():
                        if any(kw in title_lower for kw in domain_info["keywords"]):
                            matched_domain = domain_info["name_ko"]
                            break

                    if not matched_domain:
                        matched_domain = project_primary_domain or "기타"

                    if matched_domain not in domain_data:
                        domain_data[matched_domain] = {
                            "technologies": [],
                            "implementations": [],
                            "databases": [],
                            "_tech_set": set(),
                            "_impl_set": set(),
                            "_db_set": set(),
                        }

                    for item in items:
                        if isinstance(item, str):
                            stripped = item.strip()
                            if stripped:
                                domain_data[matched_domain]["_impl_set"].add(stripped)

            # Add key_tasks to implementations
            for task in key_tasks:
                task_lower = task.lower()
                task_assigned = False

                for domain_key, domain_info in DOMAIN_CATEGORIES.items():
                    if any(kw in task_lower for kw in domain_info["keywords"]):
                        domain_name = domain_info["name_ko"]
                        if domain_name not in domain_data:
                            domain_data[domain_name] = {
                                "technologies": [],
                                "implementations": [],
                                "databases": [],
                                "_tech_set": set(),
                                "_impl_set": set(),
                                "_db_set": set(),
                            }
                        domain_data[domain_name]["_impl_set"].add(task)
                        task_assigned = True
                        break

                if not task_assigned and project_primary_domain:
                    domain_data[project_primary_domain]["_impl_set"].add(task)

        # Convert sets to lists and format technologies
        for domain_name in domain_data:
            techs = sorted(domain_data[domain_name]["_tech_set"])
            # Group technologies into comma-separated lines (max 4-5 per line)
            if techs:
                domain_data[domain_name]["technologies"] = [", ".join(techs)]

            domain_data[domain_name]["implementations"] = list(domain_data[domain_name]["_impl_set"])[:10]

            dbs = sorted(domain_data[domain_name]["_db_set"])
            if dbs:
                domain_data[domain_name]["databases"] = [", ".join(dbs)]

            # Clean up internal sets
            del domain_data[domain_name]["_tech_set"]
            del domain_data[domain_name]["_impl_set"]
            del domain_data[domain_name]["_db_set"]

        # Sort by priority
        priority_order = ["Backend", "AI/ML", "Frontend", "Mobile/Cross-Platform", "데이터 분석", "Database", "DevOps/인프라", "IoT/임베디드", "기타"]
        sorted_data = {}
        for domain in priority_order:
            if domain in domain_data:
                sorted_data[domain] = domain_data[domain]
        for domain in domain_data:
            if domain not in sorted_data:
                sorted_data[domain] = domain_data[domain]

        return sorted_data
