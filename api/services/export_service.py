"""
Export Service - Generate project reports in various formats
Based on portfolio project's report styles (PROJECT_PERFORMANCE_SUMMARY.md)
"""

from typing import List, Dict, Any, Optional
from datetime import date, datetime
from pathlib import Path
import os
import re

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from api.models.user import User
from api.models.company import Company
from api.models.project import Project, ProjectTechnology
from api.models.achievement import ProjectAchievement
from api.models.repo_analysis import RepoAnalysis
from api.models.repo_analysis_edits import RepoAnalysisEdits
from api.config import get_settings

settings = get_settings()


class ExportService:
    """Service for exporting project reports to Markdown and Word formats"""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.result_dir = settings.result_dir

    def _format_date(self, d: Optional[date]) -> str:
        """Format date as YYYY.MM or YYYY-MM-DD"""
        if not d:
            return "?"
        return d.strftime("%Y.%m")

    def _format_date_range(self, start: Optional[date], end: Optional[date], is_current: bool = False) -> str:
        """Format date range"""
        start_str = self._format_date(start)
        if is_current or end is None:
            end_str = "진행중"
        else:
            end_str = self._format_date(end)
        return f"{start_str} ~ {end_str}"

    async def _get_analyzed_projects(self, user_id: int) -> List[Dict[str, Any]]:
        """Get all analyzed projects with their data"""
        # Get user
        user_result = await self.db.execute(
            select(User).where(User.id == user_id)
        )
        user = user_result.scalar_one_or_none()
        if not user:
            raise ValueError(f"User {user_id} not found")

        # Get companies
        companies_result = await self.db.execute(
            select(Company).where(Company.user_id == user_id)
        )
        companies = {c.id: c for c in companies_result.scalars().all()}

        # Get analyzed projects
        projects_result = await self.db.execute(
            select(Project)
            .where(Project.user_id == user_id, Project.is_analyzed == True)
            .options(
                selectinload(Project.technologies).selectinload(ProjectTechnology.technology),
                selectinload(Project.achievements)
            )
            .order_by(Project.start_date.desc())
        )
        projects = projects_result.scalars().all()

        result = []
        for project in projects:
            # Get repo analysis
            analysis_result = await self.db.execute(
                select(RepoAnalysis).where(RepoAnalysis.project_id == project.id)
            )
            analysis = analysis_result.scalar_one_or_none()

            # Get user edits if analysis exists
            edits = None
            if analysis:
                edits_result = await self.db.execute(
                    select(RepoAnalysisEdits).where(RepoAnalysisEdits.repo_analysis_id == analysis.id)
                )
                edits = edits_result.scalar_one_or_none()

            company = companies.get(project.company_id)

            result.append({
                "project": project,
                "analysis": analysis,
                "edits": edits,
                "company": company,
            })

        return result, user

    def _get_effective_key_tasks(self, analysis: Any, edits: Any) -> List[str]:
        """Get effective key_tasks (edited or original)"""
        if edits and edits.key_tasks_modified and edits.key_tasks is not None:
            return edits.key_tasks
        return analysis.key_tasks if analysis and analysis.key_tasks else []

    def _get_effective_implementation_details(self, analysis: Any, edits: Any) -> List[str]:
        """Get effective implementation_details (edited or original)"""
        if edits and edits.implementation_details_modified and edits.implementation_details is not None:
            return edits.implementation_details
        return analysis.implementation_details if analysis and analysis.implementation_details else []

    def _get_effective_detailed_achievements(self, analysis: Any, edits: Any) -> Dict:
        """Get effective detailed_achievements (edited or original)"""
        if edits and edits.detailed_achievements_modified and edits.detailed_achievements is not None:
            return edits.detailed_achievements
        return analysis.detailed_achievements if analysis and analysis.detailed_achievements else {}

    def _generate_toc(self, projects_data: List[Dict]) -> str:
        """Generate table of contents"""
        lines = ["## 목차", ""]
        lines.append("1. [전체 프로젝트 개요](#전체-프로젝트-개요)")

        for idx, data in enumerate(projects_data, 2):
            project = data["project"]
            anchor = self._make_anchor(project.name)
            lines.append(f"{idx}. [{project.name}](#{anchor})")

        lines.append("")
        return "\n".join(lines)

    def _make_anchor(self, name: str) -> str:
        """Make GitHub-style anchor from project name"""
        anchor = name.lower()
        anchor = re.sub(r'[^\w\s가-힣-]', '', anchor)
        anchor = re.sub(r'\s+', '-', anchor)
        return anchor

    def _generate_overview_table(self, projects_data: List[Dict], user: Any) -> str:
        """Generate overview statistics table"""
        total_commits = sum(
            data["analysis"].total_commits if data["analysis"] else 0
            for data in projects_data
        )

        # Categorize projects by technology
        categories = {
            "Backend 시스템": [],
            "Frontend": [],
            "Mobile": [],
            "AI/ML": [],
            "IoT/하드웨어": [],
            "기타": [],
        }

        backend_techs = {"FastAPI", "Django", "Flask", "Spring", "Spring Boot", "Express", "NestJS", "Node.js"}
        frontend_techs = {"React", "Vue", "Angular", "Next.js", "Nuxt.js", "Svelte"}
        mobile_techs = {"Flutter", "React Native", "Swift", "Kotlin", "Android", "iOS"}
        ai_techs = {"TensorFlow", "PyTorch", "scikit-learn", "LSTM", "Machine Learning", "Deep Learning", "OpenCV"}
        iot_techs = {"Arduino", "Raspberry Pi", "MQTT", "BLE", "Sensor", "IoT", "Embedded"}

        for data in projects_data:
            project = data["project"]
            techs = set()
            if project.technologies:
                techs = {pt.technology.name for pt in project.technologies if pt.technology}

            categorized = False
            if techs & backend_techs:
                categories["Backend 시스템"].append(project.name)
                categorized = True
            if techs & frontend_techs:
                categories["Frontend"].append(project.name)
                categorized = True
            if techs & mobile_techs:
                categories["Mobile"].append(project.name)
                categorized = True
            if techs & ai_techs:
                categories["AI/ML"].append(project.name)
                categorized = True
            if techs & iot_techs:
                categories["IoT/하드웨어"].append(project.name)
                categorized = True
            if not categorized:
                categories["기타"].append(project.name)

        lines = [
            "## 전체 프로젝트 개요",
            "",
            "### 프로젝트 포트폴리오 통계",
            "",
            "| 분류 | 프로젝트 수 | 주요 기술 스택 |",
            "|------|------------|--------------|",
        ]

        for cat_name, cat_projects in categories.items():
            if cat_projects:
                # Get main technologies for this category
                main_techs = []
                for data in projects_data:
                    if data["project"].name in cat_projects:
                        if data["project"].technologies:
                            main_techs.extend([pt.technology.name for pt in data["project"].technologies if pt.technology])
                main_techs = list(set(main_techs))[:3]
                tech_str = ", ".join(main_techs) if main_techs else "-"
                lines.append(f"| **{cat_name}** | {len(cat_projects)} | {tech_str} |")

        lines.append(f"| **총계** | **{len(projects_data)}** | **다양한 기술 스택** |")
        lines.append("")
        return "\n".join(lines)

    def _generate_project_section_summary(self, data: Dict, idx: int) -> str:
        """Generate a summary project section with key tasks and achievements (요약 형식)"""
        project = data["project"]
        analysis = data["analysis"]
        edits = data.get("edits")
        company = data["company"]

        lines = [
            "---",
            "",
            f"## {idx}. {project.name}",
        ]

        # Key tasks section
        key_tasks = self._get_effective_key_tasks(analysis, edits)
        implementation_details = self._get_effective_implementation_details(analysis, edits)

        # Combine key_tasks and implementation_details (prefer key_tasks if both exist)
        all_tasks = key_tasks if isinstance(key_tasks, list) and key_tasks else (
            implementation_details if isinstance(implementation_details, list) else []
        )

        if all_tasks:
            lines.append("### 1. 주요 수행 업무")
            for i, task in enumerate(all_tasks, 1):
                # Handle case where task might be a dict (from implementation_details structured format)
                if isinstance(task, dict):
                    task_str = task.get("title", str(task))
                else:
                    task_str = str(task)
                lines.append(f"({i}) {task_str}")
            lines.append("")

        # Achievements section
        detailed_achievements = self._get_effective_detailed_achievements(analysis, edits)
        achievements = project.achievements if project.achievements else []

        has_achievements = (detailed_achievements and isinstance(detailed_achievements, dict)) or bool(achievements)
        if has_achievements:
            lines.append("### 2. 프로젝트 성과")
            lines.append("")

            # From detailed_achievements (LLM-generated)
            if detailed_achievements and isinstance(detailed_achievements, dict):
                for category, items in detailed_achievements.items():
                    lines.append(f"#### {category}")
                    if isinstance(items, list):
                        for item in items:
                            if isinstance(item, dict):
                                title = item.get("title", "")
                                description = item.get("description", "")
                                before = item.get("before", "")
                                after = item.get("after", "")

                                lines.append(f"**[ {title} ]**")
                                if description:
                                    lines.append(f"- {description}")
                                if before and after:
                                    lines.append(f"  - 기존: {before}")
                                    lines.append(f"  - 개선: {after}")
                            else:
                                lines.append(f"- {item}")
                    lines.append("")

            # From achievements table (if detailed_achievements is empty)
            elif achievements:
                # Group by category
                by_category = {}
                for ach in achievements:
                    cat = ach.category or "기타 성과"
                    if cat not in by_category:
                        by_category[cat] = []
                    by_category[cat].append(ach)

                for category, achs in by_category.items():
                    lines.append(f"#### {category}")
                    for ach in achs:
                        if ach.metric_value:
                            lines.append(f"- **{ach.metric_name}**: {ach.metric_value}")
                        else:
                            lines.append(f"- **{ach.metric_name}**")
                        if ach.description:
                            lines.append(f"  - {ach.description}")
                    lines.append("")

        # Code statistics (if available)
        if analysis:
            lines.append("### 3. 코드 기여 통계")
            lines.append("")
            lines.append(f"- **총 커밋**: {analysis.total_commits or 0}개")
            lines.append(f"- **사용자 커밋**: {analysis.user_commits or 0}개")
            if analysis.total_commits and analysis.total_commits > 0:
                contribution = round((analysis.user_commits or 0) / analysis.total_commits * 100, 1)
                lines.append(f"- **기여율**: {contribution}%")
            lines.append(f"- **추가된 라인**: {analysis.lines_added or 0}")
            lines.append(f"- **삭제된 라인**: {analysis.lines_deleted or 0}")
            lines.append("")

        lines.append("")
        return "\n".join(lines)

    def _generate_project_section_detailed(self, data: Dict, idx: int) -> str:
        """Generate a detailed project section (상세 형식) - DETAILED_COMPLETION_REPORT style"""
        project = data["project"]
        analysis = data["analysis"]
        edits = data.get("edits")
        company = data["company"]

        lines = [
            "---",
            "",
            f"## {idx}. {project.name}",
        ]

        # Project Overview
        lines.append("### 프로젝트 개요")
        if project.git_url:
            repo_name = project.git_url.split("/")[-1].replace(".git", "") if project.git_url else project.name
            lines.append(f"- **저장소**: {repo_name}")
            lines.append(f"- **GitHub**: {project.git_url}")
        if analysis:
            lines.append(f"- **커밋**: {analysis.total_commits or 0}개")
        lines.append(f"- **기간**: {self._format_date_range(project.start_date, project.end_date)}")
        if analysis:
            lines.append(f"- **코드 변경량**: {analysis.lines_added or 0:,} 라인 추가, {analysis.lines_deleted or 0:,} 라인 삭제")
        if project.description:
            lines.append(f"- **프로젝트 성격**: {project.description[:100]}")
        lines.append("")

        # Technology Stack with versions
        lines.append("### 기술 스택")
        if project.technologies:
            # Group by category
            tech_by_category = {}
            for pt in project.technologies:
                if pt.technology:
                    category = pt.technology.category or "기타"
                    if category not in tech_by_category:
                        tech_by_category[category] = []
                    tech_by_category[category].append(pt.technology.name)

            for category, techs in tech_by_category.items():
                techs_str = ", ".join(techs)
                lines.append(f"- **{category}**: {techs_str}")
        else:
            lines.append("- 기술 스택 정보 없음")
        lines.append("")

        # Implementation Details (주요 구현 기능)
        implementation_details = self._get_effective_implementation_details(analysis, edits)
        if implementation_details:
            lines.append("### 주요 구현 기능")
            lines.append("")
            if isinstance(implementation_details, list) and implementation_details:
                # Check if first element is a dict (structured format)
                first_elem = implementation_details[0]
                if isinstance(first_elem, dict) and "title" in first_elem:
                    # Structured format: [{"title": "...", "items": [...]}]
                    for detail in implementation_details:
                        if isinstance(detail, dict):
                            title = detail.get("title", "")
                            items = detail.get("items", [])
                            lines.append(f"#### {title}")
                            if isinstance(items, list):
                                for item in items:
                                    lines.append(f"- {item}")
                            lines.append("")
                        else:
                            # Fallback for non-dict items
                            lines.append(f"- {detail}")
                else:
                    # Simple list format
                    for i, detail in enumerate(implementation_details, 1):
                        lines.append(f"#### {i}. {detail}")
            lines.append("")

        # Development Timeline (from commit analysis)
        if analysis and analysis.commit_messages_summary:
            lines.append("### 개발 타임라인")
            lines.append("")
            # Parse commit summary for timeline info
            summary = analysis.commit_messages_summary
            try:
                if isinstance(summary, dict) and "timeline" in summary:
                    timeline = summary["timeline"]
                    if isinstance(timeline, dict):
                        for period, desc in timeline.items():
                            lines.append(f"**{period}**")
                            if isinstance(desc, list):
                                for item in desc:
                                    lines.append(f"- {item}")
                            else:
                                lines.append(f"- {desc}")
                            lines.append("")
                elif isinstance(summary, str):
                    lines.append(summary)
                    lines.append("")
            except Exception:
                # Skip timeline section if parsing fails
                pass

        # Achievements
        detailed_achievements = self._get_effective_detailed_achievements(analysis, edits)
        achievements = project.achievements if project.achievements else []

        if detailed_achievements or achievements:
            lines.append("### 주요 성과")
            lines.append("")

            if detailed_achievements and isinstance(detailed_achievements, dict):
                for category, items in detailed_achievements.items():
                    lines.append(f"#### {category}")
                    if isinstance(items, list):
                        for item in items:
                            if isinstance(item, dict):
                                title = item.get("title", "")
                                description = item.get("description", "")
                                before = item.get("before", "")
                                after = item.get("after", "")

                                lines.append(f"- **{title}**")
                                if description:
                                    lines.append(f"  - {description}")
                                if before and after:
                                    lines.append(f"  - 기존: {before} → 개선: {after}")
                            else:
                                lines.append(f"- {item}")
                    lines.append("")
            elif achievements:
                by_category = {}
                for ach in achievements:
                    cat = ach.category or "기타 성과"
                    if cat not in by_category:
                        by_category[cat] = []
                    by_category[cat].append(ach)

                for category, achs in by_category.items():
                    lines.append(f"#### {category}")
                    for ach in achs:
                        if ach.metric_value:
                            lines.append(f"- **{ach.metric_name}**: {ach.metric_value}")
                        else:
                            lines.append(f"- **{ach.metric_name}**")
                        if ach.description:
                            lines.append(f"  - {ach.description}")
                    lines.append("")

        lines.append("")
        return "\n".join(lines)

    def _generate_project_section_final(self, data: Dict, idx: int) -> str:
        """Generate a final project section (상세 요약 형식) - FINAL_PROJECT_REPORT style"""
        project = data["project"]
        analysis = data["analysis"]
        edits = data.get("edits")
        company = data["company"]

        lines = [
            "---",
            "",
            f"## {idx}. {project.name}",
        ]

        # Compact project overview
        lines.append("### 프로젝트 개요")
        if project.git_url:
            lines.append(f"- **GitHub**: {project.git_url}")
        lines.append(f"- **기간**: {self._format_date_range(project.start_date, project.end_date)}")
        if analysis:
            lines.append(f"- **커밋**: {analysis.total_commits or 0}개 | **코드 변경**: +{analysis.lines_added or 0:,} / -{analysis.lines_deleted or 0:,}")
        if project.description:
            lines.append(f"- **소개**: {project.description[:150]}")
        lines.append("")

        # Technology Stack (compact)
        if project.technologies:
            tech_names = [pt.technology.name for pt in project.technologies if pt.technology][:10]
            if tech_names:
                lines.append(f"**기술 스택**: {', '.join(tech_names)}")
                lines.append("")

        # Key Implementation Features (구현 기능 - 간략)
        implementation_details = self._get_effective_implementation_details(analysis, edits)
        if implementation_details:
            lines.append("### 주요 구현 기능")
            if isinstance(implementation_details, list) and implementation_details:
                first_elem = implementation_details[0]
                if isinstance(first_elem, dict) and "title" in first_elem:
                    for detail in implementation_details[:3]:  # Limit to top 3
                        if isinstance(detail, dict):
                            title = detail.get("title", "")
                            items = detail.get("items", [])
                            if isinstance(items, list):
                                items = items[:3]  # Limit items
                            else:
                                items = []
                            lines.append(f"**{title}**")
                            for item in items:
                                lines.append(f"- {item}")
                        else:
                            lines.append(f"- {detail}")
                else:
                    for i, detail in enumerate(implementation_details[:5], 1):  # Limit to top 5
                        lines.append(f"- {detail}")
            lines.append("")

        # Key Tasks (주요 수행 업무)
        key_tasks = self._get_effective_key_tasks(analysis, edits)
        if key_tasks and isinstance(key_tasks, list):
            lines.append("### 주요 수행 업무")
            for i, task in enumerate(key_tasks[:8], 1):  # Limit to 8
                lines.append(f"({i}) {task}")
            lines.append("")

        # Achievements (compact)
        detailed_achievements = self._get_effective_detailed_achievements(analysis, edits)
        achievements = project.achievements if project.achievements else []

        if detailed_achievements or achievements:
            lines.append("### 성과")
            if detailed_achievements and isinstance(detailed_achievements, dict):
                for category, items in list(detailed_achievements.items())[:3]:  # Limit categories
                    lines.append(f"**{category}**")
                    if isinstance(items, list):
                        for item in items[:2]:  # Limit items per category
                            if isinstance(item, dict):
                                title = item.get("title", "")
                                description = item.get("description", "")
                                lines.append(f"- {title}: {description}" if description else f"- {title}")
                            else:
                                lines.append(f"- {item}")
                    lines.append("")
            elif achievements:
                for ach in achievements[:5]:
                    if ach.metric_value:
                        lines.append(f"- **{ach.metric_name}**: {ach.metric_value}")
                    else:
                        lines.append(f"- **{ach.metric_name}**")
                lines.append("")

        lines.append("")
        return "\n".join(lines)

    def _generate_report_header(self, projects_data: List[Dict], user: Any, report_title: str) -> List[str]:
        """Generate common report header"""
        total_commits = sum(
            data["analysis"].total_commits if data["analysis"] else 0
            for data in projects_data
        )

        start_dates = [
            data["project"].start_date
            for data in projects_data
            if data["project"].start_date
        ]
        end_dates = [
            data["project"].end_date
            for data in projects_data
            if data["project"].end_date
        ]

        earliest = min(start_dates) if start_dates else None
        latest = max(end_dates) if end_dates else None

        lines = [
            f"# {report_title}",
            "",
            f"**작성일**: {datetime.now().strftime('%Y-%m-%d')}",
            f"**분석 대상**: {len(projects_data)}개 프로젝트",
            f"**총 커밋 수**: {total_commits:,}개",
        ]

        if earliest:
            period = f"{earliest.strftime('%Y-%m-%d')} ~ "
            if latest:
                period += latest.strftime('%Y-%m-%d')
            else:
                period += "현재"
            lines.append(f"**작업 기간**: {period}")

        if user.github_username:
            lines.append(f"**분석 대상자**: {user.github_username}")

        lines.extend(["", "---", ""])
        return lines

    async def generate_summary_report(self, user_id: int) -> str:
        """
        Generate summary report (요약) - PROJECT_PERFORMANCE_SUMMARY.md style

        Format:
        - Header with metadata
        - Table of contents
        - Overview statistics
        - Per-project sections with:
          - 주요 수행 업무 (key tasks) in (1), (2), (3) format
          - 프로젝트 성과 (achievements by category)
          - 코드 기여 통계
        """
        projects_data, user = await self._get_analyzed_projects(user_id)

        if not projects_data:
            return "# 프로젝트 완료결과보고서\n\n분석된 프로젝트가 없습니다."

        lines = self._generate_report_header(projects_data, user, "프로젝트 완료결과보고서")

        # Add TOC
        lines.append(self._generate_toc(projects_data))
        lines.append("---")
        lines.append("")

        # Add overview
        lines.append(self._generate_overview_table(projects_data, user))
        lines.append("---")
        lines.append("")

        # Add project sections (summary format)
        for idx, data in enumerate(projects_data, 1):
            lines.append(self._generate_project_section_summary(data, idx))

        return "\n".join(lines)

    async def generate_detailed_report(self, user_id: int) -> str:
        """
        Generate detailed report (상세) - DETAILED_COMPLETION_REPORT style

        Format:
        - Header with metadata
        - Table of contents
        - Overview statistics
        - Per-project sections with:
          - 프로젝트 개요 (저장소, GitHub, 커밋, 기간, 코드 변경량)
          - 기술 스택 (카테고리별)
          - 주요 구현 기능 (상세)
          - 개발 타임라인
          - 주요 성과 (카테고리별)
        """
        projects_data, user = await self._get_analyzed_projects(user_id)

        if not projects_data:
            return "# 프로젝트 상세 보고서\n\n분석된 프로젝트가 없습니다."

        lines = self._generate_report_header(projects_data, user, "프로젝트 상세 보고서")

        # Add TOC
        lines.append(self._generate_toc(projects_data))
        lines.append("---")
        lines.append("")

        # Add overview
        lines.append(self._generate_overview_table(projects_data, user))
        lines.append("---")
        lines.append("")

        # Add project sections (detailed format)
        for idx, data in enumerate(projects_data, 1):
            lines.append(self._generate_project_section_detailed(data, idx))

        return "\n".join(lines)

    async def generate_final_report(self, user_id: int) -> str:
        """
        Generate final report (상세 요약) - FINAL_PROJECT_REPORT style

        Format:
        - Header with metadata
        - Table of contents
        - Overview statistics
        - Per-project sections with:
          - 프로젝트 개요 (간략)
          - 기술 스택 (한 줄)
          - 주요 구현 기능 (상위 3개)
          - 주요 수행 업무 (상위 8개)
          - 성과 (상위 카테고리)
        """
        projects_data, user = await self._get_analyzed_projects(user_id)

        if not projects_data:
            return "# 프로젝트 최종 보고서\n\n분석된 프로젝트가 없습니다."

        lines = self._generate_report_header(projects_data, user, "프로젝트 최종 보고서")

        # Add TOC
        lines.append(self._generate_toc(projects_data))
        lines.append("---")
        lines.append("")

        # Add overview
        lines.append(self._generate_overview_table(projects_data, user))
        lines.append("---")
        lines.append("")

        # Add project sections (final format - concise)
        for idx, data in enumerate(projects_data, 1):
            lines.append(self._generate_project_section_final(data, idx))

        return "\n".join(lines)

    # Alias for backward compatibility
    async def generate_performance_summary_report(self, user_id: int) -> str:
        """Alias for generate_summary_report (backward compatibility)"""
        return await self.generate_summary_report(user_id)

    async def export_to_markdown(self, user_id: int, report_type: str = "summary") -> tuple[str, str]:
        """
        Export report to Markdown file

        Args:
            user_id: User ID
            report_type: "detailed", "final", or "summary"

        Returns:
            Tuple of (file_path, content)
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        if report_type == "detailed":
            content = await self.generate_detailed_report(user_id)
            filename = f"PROJECT_DETAILED_REPORT_{timestamp}.md"
        elif report_type == "final":
            content = await self.generate_final_report(user_id)
            filename = f"PROJECT_FINAL_REPORT_{timestamp}.md"
        else:  # summary (default)
            content = await self.generate_summary_report(user_id)
            filename = f"PROJECT_SUMMARY_REPORT_{timestamp}.md"

        os.makedirs(self.result_dir, exist_ok=True)
        file_path = self.result_dir / filename

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return str(file_path), content

    async def export_to_docx(self, user_id: int, report_type: str = "summary") -> str:
        """
        Export report to Word document

        Args:
            user_id: User ID
            report_type: "detailed", "final", or "summary"

        Returns:
            File path to generated DOCX
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        if report_type == "detailed":
            content = await self.generate_detailed_report(user_id)
            filename = f"PROJECT_DETAILED_REPORT_{timestamp}.docx"
        elif report_type == "final":
            content = await self.generate_final_report(user_id)
            filename = f"PROJECT_FINAL_REPORT_{timestamp}.docx"
        else:  # summary (default)
            content = await self.generate_summary_report(user_id)
            filename = f"PROJECT_SUMMARY_REPORT_{timestamp}.docx"

        os.makedirs(self.result_dir, exist_ok=True)
        file_path = self.result_dir / filename

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(11)

        # Parse markdown and convert to DOCX
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()

            if stripped.startswith('# '):
                para = doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                para = doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                para = doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('#### '):
                para = doc.add_paragraph()
                run = para.add_run(stripped[5:])
                run.bold = True
            elif stripped.startswith('- **') and '**:' in stripped:
                # Key-value pair like "- **기간**: 2024.01 ~ 2024.06"
                para = doc.add_paragraph(style='List Bullet')
                # Parse bold text
                match = re.match(r'- \*\*(.+?)\*\*:?\s*(.*)', stripped)
                if match:
                    run = para.add_run(match.group(1))
                    run.bold = True
                    if match.group(2):
                        para.add_run(f": {match.group(2)}")
                else:
                    para.add_run(stripped[2:])
            elif stripped.startswith('- '):
                para = doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped.startswith('(') and ')' in stripped[:5]:
                # Numbered item like "(1) 내용"
                para = doc.add_paragraph(stripped)
            elif stripped.startswith('**[') and ']**' in stripped:
                # Bold section title like "**[ Export 최적화 ]**"
                para = doc.add_paragraph()
                title = re.search(r'\*\*\[\s*(.+?)\s*\]\*\*', stripped)
                if title:
                    run = para.add_run(f"[ {title.group(1)} ]")
                    run.bold = True
                else:
                    run = para.add_run(stripped.replace('**', ''))
                    run.bold = True
            elif stripped.startswith('**') and stripped.endswith('**'):
                para = doc.add_paragraph()
                run = para.add_run(stripped[2:-2])
                run.bold = True
            elif stripped.startswith('|'):
                # Table row - simplified handling
                para = doc.add_paragraph(stripped)
                para.style = 'Normal'
            elif stripped == '---':
                doc.add_paragraph()  # Add spacing
            elif stripped:
                # Handle inline bold
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.+?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        para.add_run(part)

        doc.save(file_path)
        return str(file_path)

    async def get_export_preview(self, user_id: int, report_type: str = "summary") -> Dict[str, Any]:
        """
        Get preview data for export

        Args:
            user_id: User ID
            report_type: "detailed", "final", or "summary"

        Returns:
            Dict with project count, preview content, etc.
        """
        projects_data, user = await self._get_analyzed_projects(user_id)

        if report_type == "detailed":
            content = await self.generate_detailed_report(user_id)
        elif report_type == "final":
            content = await self.generate_final_report(user_id)
        else:  # summary (default)
            content = await self.generate_summary_report(user_id)

        # Calculate stats
        total_commits = sum(
            data["analysis"].total_commits if data["analysis"] else 0
            for data in projects_data
        )

        has_key_tasks = any(
            self._get_effective_key_tasks(data["analysis"], data.get("edits"))
            for data in projects_data
        )

        has_achievements = any(
            self._get_effective_detailed_achievements(data["analysis"], data.get("edits")) or
            (data["project"].achievements and len(data["project"].achievements) > 0)
            for data in projects_data
        )

        return {
            "project_count": len(projects_data),
            "total_commits": total_commits,
            "has_key_tasks": has_key_tasks,
            "has_achievements": has_achievements,
            "preview": content[:2000] + "..." if len(content) > 2000 else content,
            "full_content": content,
        }

    # ==================== Single Project Export Methods ====================

    async def _get_single_project(self, project_id: int) -> Dict[str, Any]:
        """Get a single project with its data for export"""
        # Get project with technologies and achievements
        project_result = await self.db.execute(
            select(Project)
            .where(Project.id == project_id)
            .options(
                selectinload(Project.technologies).selectinload(ProjectTechnology.technology),
                selectinload(Project.achievements)
            )
        )
        project = project_result.scalar_one_or_none()
        if not project:
            raise ValueError(f"Project {project_id} not found")

        # Get repo analysis
        analysis_result = await self.db.execute(
            select(RepoAnalysis).where(RepoAnalysis.project_id == project_id)
        )
        analysis = analysis_result.scalar_one_or_none()

        # Get user edits if analysis exists
        edits = None
        if analysis:
            edits_result = await self.db.execute(
                select(RepoAnalysisEdits).where(RepoAnalysisEdits.repo_analysis_id == analysis.id)
            )
            edits = edits_result.scalar_one_or_none()

        # Get company if exists
        company = None
        if project.company_id:
            company_result = await self.db.execute(
                select(Company).where(Company.id == project.company_id)
            )
            company = company_result.scalar_one_or_none()

        # Get user
        user_result = await self.db.execute(
            select(User).where(User.id == project.user_id)
        )
        user = user_result.scalar_one_or_none()

        return {
            "project": project,
            "analysis": analysis,
            "edits": edits,
            "company": company,
            "user": user,
        }

    async def generate_single_project_report(
        self,
        project_id: int,
        report_type: str = "summary"
    ) -> str:
        """
        Generate a report for a single project

        Args:
            project_id: Project ID
            report_type: "detailed", "final", or "summary"

        Returns:
            Markdown content
        """
        data = await self._get_single_project(project_id)
        project = data["project"]
        user = data["user"]

        # Build header
        lines = [
            f"# {project.name}",
            "",
            f"**작성일**: {datetime.now().strftime('%Y-%m-%d')}",
        ]

        if user and user.github_username:
            lines.append(f"**분석 대상자**: {user.github_username}")

        lines.extend(["", "---", ""])

        # Add project section based on report type
        if report_type == "detailed":
            lines.append(self._generate_project_section_detailed(data, 1).replace("## 1. ", "## "))
        elif report_type == "final":
            lines.append(self._generate_project_section_final(data, 1).replace("## 1. ", "## "))
        else:  # summary
            lines.append(self._generate_project_section_summary(data, 1).replace("## 1. ", "## "))

        return "\n".join(lines)

    async def export_single_project_to_markdown(
        self,
        project_id: int,
        report_type: str = "summary"
    ) -> tuple[str, str]:
        """
        Export single project report to Markdown file

        Args:
            project_id: Project ID
            report_type: "detailed", "final", or "summary"

        Returns:
            Tuple of (file_path, content)
        """
        data = await self._get_single_project(project_id)
        project = data["project"]

        content = await self.generate_single_project_report(project_id, report_type)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_name = re.sub(r'[^\w\s-]', '', project.name).strip().replace(' ', '_')[:30]
        filename = f"{safe_name}_{report_type}_{timestamp}.md"

        os.makedirs(self.result_dir, exist_ok=True)
        file_path = self.result_dir / filename

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return str(file_path), content

    async def export_single_project_to_docx(
        self,
        project_id: int,
        report_type: str = "summary"
    ) -> str:
        """
        Export single project report to Word document

        Args:
            project_id: Project ID
            report_type: "detailed", "final", or "summary"

        Returns:
            File path to generated DOCX
        """
        from docx import Document
        from docx.shared import Pt

        data = await self._get_single_project(project_id)
        project = data["project"]

        content = await self.generate_single_project_report(project_id, report_type)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_name = re.sub(r'[^\w\s-]', '', project.name).strip().replace(' ', '_')[:30]
        filename = f"{safe_name}_{report_type}_{timestamp}.docx"

        os.makedirs(self.result_dir, exist_ok=True)
        file_path = self.result_dir / filename

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(11)

        # Parse markdown and convert to DOCX (reuse existing logic)
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()

            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('#### '):
                para = doc.add_paragraph()
                run = para.add_run(stripped[5:])
                run.bold = True
            elif stripped.startswith('- **') and '**:' in stripped:
                para = doc.add_paragraph(style='List Bullet')
                match = re.match(r'- \*\*(.+?)\*\*:?\s*(.*)', stripped)
                if match:
                    run = para.add_run(match.group(1))
                    run.bold = True
                    if match.group(2):
                        para.add_run(f": {match.group(2)}")
                else:
                    para.add_run(stripped[2:])
            elif stripped.startswith('- '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped.startswith('(') and ')' in stripped[:5]:
                doc.add_paragraph(stripped)
            elif stripped.startswith('**[') and ']**' in stripped:
                para = doc.add_paragraph()
                title = re.search(r'\*\*\[\s*(.+?)\s*\]\*\*', stripped)
                if title:
                    run = para.add_run(f"[ {title.group(1)} ]")
                    run.bold = True
                else:
                    run = para.add_run(stripped.replace('**', ''))
                    run.bold = True
            elif stripped.startswith('**') and stripped.endswith('**'):
                para = doc.add_paragraph()
                run = para.add_run(stripped[2:-2])
                run.bold = True
            elif stripped == '---':
                doc.add_paragraph()
            elif stripped:
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.+?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        para.add_run(part)

        doc.save(file_path)
        return str(file_path)

    async def get_single_project_preview(
        self,
        project_id: int,
        report_type: str = "summary"
    ) -> Dict[str, Any]:
        """
        Get preview data for single project export

        Args:
            project_id: Project ID
            report_type: "detailed", "final", or "summary"

        Returns:
            Dict with project info, preview content, etc.
        """
        data = await self._get_single_project(project_id)
        project = data["project"]
        analysis = data["analysis"]
        edits = data.get("edits")

        content = await self.generate_single_project_report(project_id, report_type)

        total_commits = analysis.total_commits if analysis else 0
        has_key_tasks = bool(self._get_effective_key_tasks(analysis, edits))
        has_achievements = bool(
            self._get_effective_detailed_achievements(analysis, edits) or
            (project.achievements and len(project.achievements) > 0)
        )

        return {
            "project_id": project_id,
            "project_name": project.name,
            "total_commits": total_commits,
            "has_key_tasks": has_key_tasks,
            "has_achievements": has_achievements,
            "preview": content[:2000] + "..." if len(content) > 2000 else content,
            "full_content": content,
        }
