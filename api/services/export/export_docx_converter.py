"""
DOCX Converter - Markdown to DOCX conversion utilities.

Extracted from export_service.py for better modularity.
Contains helper function for converting Markdown content to DOCX.
"""
import re
from typing import Any

from docx import Document

from api.services.docx.docx_styles import DocxStyler


def convert_markdown_to_docx(
    content: str,
    styler: DocxStyler,
) -> Document:
    """
    Convert Markdown content to a Word document with proper styling.

    Uses DocxStyler for consistent formatting:
    - Title (대제목): 24pt, bold, black
    - Heading1 (중제목): 18pt, bold, black
    - Heading2 (소제목): 14pt, bold, black
    - Normal text: 11pt, black

    Args:
        content: Markdown content to convert
        styler: DocxStyler instance for formatting

    Returns:
        Document object
    """
    doc = Document()
    styler.setup_document(doc)

    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()

        if stripped.startswith('# '):
            # Title (대제목) - 24pt
            styler.add_title(doc, stripped[2:], center=False)
        elif stripped.startswith('## '):
            # Heading1 (중제목) - 18pt
            styler.add_heading1(doc, stripped[3:])
        elif stripped.startswith('### '):
            # Heading2 (소제목) - 14pt
            styler.add_heading2(doc, stripped[4:])
        elif stripped.startswith('#### '):
            # Heading3 (소소제목) - 12pt
            styler.add_heading3(doc, stripped[5:])
        elif stripped.startswith('- **') and '**:' in stripped:
            # Key-value pair like "- **기간**: 2024.01 ~ 2024.06"
            match = re.match(r'- \*\*(.+?)\*\*:?\s*(.*)', stripped)
            if match:
                styler.add_key_value(doc, match.group(1), match.group(2), as_bullet=True)
            else:
                styler.add_bullet(doc, stripped[2:])
        elif stripped.startswith('- '):
            styler.add_bullet(doc, stripped[2:])
        elif stripped.startswith('(') and ')' in stripped[:5]:
            # Numbered item like "(1) 내용"
            styler.add_paragraph(doc, stripped)
        elif stripped.startswith('**[') and ']**' in stripped:
            # Bold section title like "**[ Export 최적화 ]**"
            title = re.search(r'\*\*\[\s*(.+?)\s*\]\*\*', stripped)
            if title:
                styler.add_section_title(doc, f"[ {title.group(1)} ]")
            else:
                styler.add_bold_text(doc, stripped.replace('**', ''))
        elif stripped.startswith('**') and stripped.endswith('**'):
            styler.add_bold_text(doc, stripped[2:-2])
        elif stripped.startswith('|'):
            # Table row - simplified handling
            styler.add_paragraph(doc, stripped)
        elif stripped == '---':
            styler.add_spacing(doc)
        elif stripped:
            # Handle inline bold
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.+?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = styler.config.normal.pt_size
                    run.font.name = styler.config.font_name
                    run.font.color.rgb = styler.config.normal.rgb_color
                else:
                    run = para.add_run(part)
                    run.font.size = styler.config.normal.pt_size
                    run.font.name = styler.config.font_name
                    run.font.color.rgb = styler.config.normal.rgb_color

    return doc
