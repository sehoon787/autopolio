"""
Word Document Generator - Creates properly formatted DOCX files matching original templates.

This module generates Word documents with actual tables (not markdown conversion)
that match the exact format of the original Korean resume/career templates.

Original document analysis (from 김세훈 경력기술서/이력서 templates):
- Font: 나눔고딕OTF
- Title: 26pt
- Name: 18pt Bold
- Summary: 8pt
- Table content: 10pt (implied)
- Project header: 14pt Bold
- Section title: 16pt Bold (자기소개)

- Page margins: top=3.0cm, bottom/left/right=2.54cm
- 3-column table widths: [3.24cm, 3.50cm, 9.16cm]
- 2-column table widths: [2.75cm, 13.15cm]
- Project row height: first=5.20cm, others=1.95cm
"""
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime
import os

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


class DocxGenerator:
    """Generate Word documents matching original Korean resume format."""

    def __init__(self):
        # Font settings - match original templates
        self.font_title = "나눔고딕OTF"
        self.font_body = "나눔고딕OTF"
        self.font_fallback = "맑은 고딕"  # Fallback if 나눔고딕OTF not available

        # Font sizes (in points) - exact match to original
        self.size_document_title = 26
        self.size_name = 18
        self.size_summary = 8  # Original uses 8pt
        self.size_section_header = 12
        self.size_table_content = 10
        self.size_project_title = 14
        self.size_subsection_title = 16  # For 경력기술서/자기소개 headers

        # Page margins (in cm) - exact match to original
        self.margin_top = Cm(3.0)
        self.margin_bottom = Cm(2.54)
        self.margin_left = Cm(2.54)
        self.margin_right = Cm(2.54)

        # 3-column table widths (in cm) - exact match to original
        self.col_width_section = Cm(3.24)  # Section header column
        self.col_width_label = Cm(3.50)    # Label column
        self.col_width_value = Cm(9.16)    # Value column

        # 2-column table widths (in cm) - exact match to original
        self.col_width_2col_section = Cm(2.75)   # Section header
        self.col_width_2col_content = Cm(13.15)  # Content

        # Row heights (in cm)
        self.row_height_project_first = Cm(5.20)  # First row of project section
        self.row_height_project_other = Cm(1.95)  # Other rows

    def _set_cell_borders(self, cell, color: str = "auto", size: int = 12):
        """Set cell borders matching original template.

        Original uses sz=12 (0.5pt), color=auto
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing borders
        for old_border in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old_border)

        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(size))
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def _set_cell_shading(self, cell, color: str):
        """Set cell background color."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        tcPr.append(shading)

    def _merge_cells_vertically(self, table, col_idx: int, start_row: int, end_row: int):
        """Merge cells vertically in a column."""
        for row_idx in range(start_row, end_row + 1):
            cell = table.cell(row_idx, col_idx)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            if row_idx == start_row:
                # First cell - vMerge restart
                vMerge = OxmlElement('w:vMerge')
                vMerge.set(qn('w:val'), 'restart')
            else:
                # Subsequent cells - vMerge continue
                vMerge = OxmlElement('w:vMerge')
                # Clear the cell text
                cell.text = ""

            tcPr.append(vMerge)

    def _set_page_margins(self, doc: Document):
        """Set page margins and size to match original template (A4)."""
        for section in doc.sections:
            # Set page size to A4
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            # Set margins
            section.top_margin = self.margin_top
            section.bottom_margin = self.margin_bottom
            section.left_margin = self.margin_left
            section.right_margin = self.margin_right

    def _set_table_column_widths(self, table, widths: list):
        """Set table column widths properly using XML.

        python-docx cell.width doesn't always work, so we set widths via XML.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        # Remove existing width settings
        for old_width in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old_width)

        # Set table width to auto
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'auto')
        tblPr.append(tblW)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

        # Set column widths using tblGrid
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            tbl.insert(1, tblGrid)
        else:
            # Clear existing gridCol elements
            for col in list(tblGrid):
                tblGrid.remove(col)

        for width in widths:
            gridCol = OxmlElement('w:gridCol')
            # Convert cm to twips (1 cm = 567 twips)
            if isinstance(width, Cm):
                twips = int(width.cm * 567)
            else:
                twips = int(width * 567)
            gridCol.set(qn('w:w'), str(twips))
            tblGrid.append(gridCol)

        # Also set cell widths directly
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths):
                    width = widths[idx]
                    if isinstance(width, Cm):
                        cell.width = width
                    else:
                        cell.width = Cm(width)

    def _set_run_font(self, run, font_size: int, bold: bool = False):
        """Set font properties for a run."""
        run.font.name = self.font_body
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Set East Asian font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_body)

    def _set_paragraph_spacing(self, para, space_before: int = 0, space_after: int = 0):
        """Set paragraph spacing to match original template.

        Original uses: space_after=0, line=240 (single), lineRule=auto
        """
        pPr = para._p.get_or_add_pPr()

        # Set spacing
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), str(int(space_before * 20)))  # Convert pt to twips
        spacing.set(qn('w:after'), str(int(space_after * 20)))
        spacing.set(qn('w:line'), '240')  # Single line spacing
        spacing.set(qn('w:lineRule'), 'auto')

        # Remove existing spacing
        for old_spacing in pPr.findall(qn('w:spacing')):
            pPr.remove(old_spacing)
        pPr.append(spacing)

    def _add_paragraph_with_style(
        self,
        doc: Document,
        text: str,
        font_size: int,
        bold: bool = False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_after: int = 0
    ):
        """Add a paragraph with specific styling."""
        para = doc.add_paragraph()
        para.alignment = alignment

        # Set spacing to match original
        self._set_paragraph_spacing(para, space_before=0, space_after=space_after)

        run = para.add_run(text)
        self._set_run_font(run, font_size, bold)

        return para

    def _create_info_table(
        self,
        doc: Document,
        section_title: str,
        rows: List[tuple],  # List of (label, value) tuples
        title_rows: int = None,  # Number of rows to merge for title column
        value_bold: bool = False  # Whether to make value column bold
    ):
        """Create a standard info table with section header in first column.

        Args:
            doc: Document to add table to
            section_title: Title for the section (e.g., "인적사항")
            rows: List of (label, value) tuples
            title_rows: Number of rows to merge for title (default: all rows)
            value_bold: Whether to bold the value column (for company names etc.)
        """
        if not rows:
            return None

        num_rows = len(rows)
        title_rows = title_rows or num_rows

        # Create 3-column table
        table = doc.add_table(rows=num_rows, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths - exact match to original [3.24cm, 3.50cm, 9.16cm]
        self._set_table_column_widths(table, [3.24, 3.50, 9.16])

        # Fill in the table
        for row_idx, (label, value) in enumerate(rows):
            # Section header (first column)
            cell0 = table.cell(row_idx, 0)
            if row_idx == 0:
                # Clear existing paragraphs and add new one
                cell0.text = ""
                para = cell0.paragraphs[0]
                run = para.add_run(section_title)
                self._set_run_font(run, self.size_table_content, bold=True)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Label (second column)
            cell1 = table.cell(row_idx, 1)
            cell1.text = ""
            para = cell1.paragraphs[0]
            run = para.add_run(label)
            self._set_run_font(run, self.size_table_content, bold=False)
            cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Value (third column)
            cell2 = table.cell(row_idx, 2)
            cell2.text = ""
            para = cell2.paragraphs[0]
            value_text = str(value) if value else ""
            run = para.add_run(value_text)
            self._set_run_font(run, self.size_table_content, bold=value_bold)
            cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Set borders for all cells
            self._set_cell_borders(cell0)
            self._set_cell_borders(cell1)
            self._set_cell_borders(cell2)

        # Merge title cells vertically
        if num_rows > 1:
            self._merge_cells_vertically(table, 0, 0, min(title_rows - 1, num_rows - 1))

        doc.add_paragraph()  # Add spacing after table
        return table

    def _create_career_table_with_bold_company(
        self,
        doc: Document,
        section_title: str,
        rows: List[tuple]
    ):
        """Create career table with bold company name in first row's value.

        Args:
            doc: Document to add table to
            section_title: Title for the section (e.g., "경력사항")
            rows: List of (label, value) tuples - first row value will be bold
        """
        if not rows:
            return None

        num_rows = len(rows)

        # Create 3-column table
        table = doc.add_table(rows=num_rows, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths - exact match to original [3.24cm, 3.50cm, 9.16cm]
        self._set_table_column_widths(table, [3.24, 3.50, 9.16])

        # Fill in the table
        for row_idx, (label, value) in enumerate(rows):
            # Section header (first column)
            cell0 = table.cell(row_idx, 0)
            if row_idx == 0:
                cell0.text = ""
                para = cell0.paragraphs[0]
                run = para.add_run(section_title)
                self._set_run_font(run, self.size_table_content, bold=True)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Label (second column)
            cell1 = table.cell(row_idx, 1)
            cell1.text = ""
            para = cell1.paragraphs[0]
            run = para.add_run(label)
            self._set_run_font(run, self.size_table_content, bold=False)
            cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Value (third column) - first row (company name) is bold
            cell2 = table.cell(row_idx, 2)
            cell2.text = ""
            para = cell2.paragraphs[0]
            value_text = str(value) if value else ""
            is_company_row = row_idx == 0  # First row is company name
            run = para.add_run(value_text)
            self._set_run_font(run, self.size_table_content, bold=is_company_row)
            cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Set borders for all cells
            self._set_cell_borders(cell0)
            self._set_cell_borders(cell1)
            self._set_cell_borders(cell2)

        # Merge title cells vertically
        if num_rows > 1:
            self._merge_cells_vertically(table, 0, 0, num_rows - 1)

        doc.add_paragraph()  # Add spacing after table
        return table

    def _create_project_table(
        self,
        doc: Document,
        projects: List[Dict[str, Any]],
        section_title: str = "주요 프로젝트"
    ):
        """Create project section table in 2-column format matching original layout."""
        if not projects:
            return None

        # Create 2-column table - exact widths from original
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths - exact match [2.75cm, 13.15cm]
        self._set_table_column_widths(table, [2.75, 13.15])

        # First row - section header
        cell0 = table.cell(0, 0)
        cell0.text = ""
        para = cell0.paragraphs[0]
        run = para.add_run(section_title)
        self._set_run_font(run, self.size_table_content, bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Project content cell
        cell1 = table.cell(0, 1)
        cell1.text = ""

        # Build project content
        for p_idx, project in enumerate(projects):
            if p_idx > 0:
                # Add separator line between projects
                para = cell1.add_paragraph()
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)

            # Project header: company / department / position
            company_info = project.get("company_name", "")
            if project.get("department"):
                company_info += f" / {project['department']}"
            if project.get("position"):
                company_info += f" / {project['position']}"

            if company_info:
                if p_idx == 0:
                    # First project - use existing paragraph
                    para = cell1.paragraphs[0]
                else:
                    para = cell1.add_paragraph()
                run = para.add_run(company_info)
                self._set_run_font(run, self.size_project_title, bold=True)

            # Company description (회사 소개)
            if project.get("company_description"):
                para = cell1.add_paragraph()
                run = para.add_run(f"회사 소개: {project['company_description']}")
                self._set_run_font(run, self.size_table_content, bold=False)

            # Project name with period
            para = cell1.add_paragraph()
            period = f"{project.get('start_date', '')} ~ {project.get('end_date', '진행중')}"
            project_line = f"프로젝트: {project.get('name', '')} ({period})"
            run = para.add_run(project_line)
            self._set_run_font(run, self.size_table_content, bold=True)

            # Role
            if project.get("role"):
                para = cell1.add_paragraph()
                run = para.add_run(f"역할: {project['role']}")
                self._set_run_font(run, self.size_table_content, bold=False)

            # Technologies
            if project.get("technologies"):
                para = cell1.add_paragraph()
                run = para.add_run(f"기술스택: {project['technologies']}")
                self._set_run_font(run, self.size_table_content, bold=False)

            # Description
            if project.get("description"):
                para = cell1.add_paragraph()
                run = para.add_run(project['description'])
                self._set_run_font(run, self.size_table_content, bold=False)

            # Key tasks (주요 수행 업무)
            key_tasks = project.get("key_tasks", [])
            if key_tasks:
                para = cell1.add_paragraph()
                para.paragraph_format.space_before = Pt(3)
                run = para.add_run("주요 수행 업무:")
                self._set_run_font(run, self.size_table_content, bold=True)

                for task in key_tasks:
                    para = cell1.add_paragraph()
                    run = para.add_run(f"• {task}")
                    self._set_run_font(run, self.size_table_content, bold=False)

            # Achievements (성과)
            # Priority 1: achievements_detailed_list (with description) if available
            # Priority 2: achievements_summary_list (title only) - DEFAULT
            # Priority 3: achievements (old format)
            achievements_detailed = project.get("achievements_detailed_list", [])
            achievements_summary = project.get("achievements_summary_list", [])

            # Determine which list to use
            use_detailed = achievements_detailed and len(achievements_detailed) > 0
            achievements_list = achievements_detailed if use_detailed else achievements_summary

            if not achievements_list:
                # Fallback to old achievements format
                achievements_list = project.get("achievements", [])

            if achievements_list:
                para = cell1.add_paragraph()
                para.paragraph_format.space_before = Pt(3)
                run = para.add_run("성과:")
                self._set_run_font(run, self.size_table_content, bold=True)

                for ach in achievements_list:
                    if isinstance(ach, dict):
                        # Format: [category] title
                        category = ach.get("category", ach.get("metric_name", "성과"))
                        title = ach.get("title", ach.get("metric_value", ""))
                        description = ach.get("description", "")

                        # Achievement line: [category] title
                        ach_text = f"[{category}] {title}"
                        para = cell1.add_paragraph()
                        run = para.add_run(ach_text)
                        self._set_run_font(run, self.size_table_content, bold=False)

                        # Add description on new line if available (for detailed format)
                        if use_detailed and description:
                            para = cell1.add_paragraph()
                            para.paragraph_format.left_indent = Pt(12)
                            run = para.add_run(f"→ {description}")
                            self._set_run_font(run, self.size_table_content - 1, bold=False)
                    else:
                        ach_text = f"• {ach}"
                        para = cell1.add_paragraph()
                        run = para.add_run(ach_text)
                        self._set_run_font(run, self.size_table_content, bold=False)

        # Set borders for all cells
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_borders(cell)

        doc.add_paragraph()
        return table

    def _create_skills_table(
        self,
        doc: Document,
        skills_data: Dict[str, str],
        section_title: str = "기술스택"
    ):
        """Create skills table matching original format.

        Original format: 2 rows x 3 cols with widths [3.24cm, 4.51cm, 8.15cm]
        Row 1: Programming Languages / Framework -> values
        Row 2: Server Tooling/DevOps / Environment -> values
        """
        if not skills_data:
            return None

        # Build content based on available data
        prog_lang = skills_data.get("programming_languages", "")
        frameworks = skills_data.get("frameworks", "")
        databases = skills_data.get("databases", "")
        cloud = skills_data.get("cloud", "") or skills_data.get("devops", "")
        tools = skills_data.get("tools", "") or skills_data.get("tooling", "")

        # Combine values for each row
        row1_label = "Programing Languages\nFramework"
        row1_value = ""
        if prog_lang:
            row1_value = prog_lang
        if frameworks:
            if row1_value:
                row1_value += "\n" + frameworks
            else:
                row1_value = frameworks

        row2_label = "Server Tooling/DevOps\nEnvironment"
        row2_value = ""
        if cloud:
            row2_value = cloud
        if tools:
            if row2_value:
                row2_value += "\n" + tools
            else:
                row2_value = tools
        if databases:
            if row2_value:
                row2_value += "\n" + databases
            else:
                row2_value = databases

        # Skip if no data
        if not row1_value and not row2_value:
            return None

        # Create 3-column table with different widths for skills
        num_rows = 0
        if row1_value:
            num_rows += 1
        if row2_value:
            num_rows += 1

        if num_rows == 0:
            return None

        table = doc.add_table(rows=num_rows, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths for skills table [3.24cm, 4.51cm, 8.15cm]
        self._set_table_column_widths(table, [3.24, 4.51, 8.15])

        row_idx = 0

        if row1_value:
            # Row 1: Programming Languages / Framework
            cell0 = table.cell(row_idx, 0)
            if row_idx == 0:
                cell0.text = ""
                para = cell0.paragraphs[0]
                run = para.add_run(section_title)
                self._set_run_font(run, self.size_table_content, bold=True)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            cell1 = table.cell(row_idx, 1)
            cell1.text = ""
            para = cell1.paragraphs[0]
            run = para.add_run(row1_label)
            self._set_run_font(run, self.size_table_content, bold=False)
            cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            cell2 = table.cell(row_idx, 2)
            cell2.text = ""
            para = cell2.paragraphs[0]
            run = para.add_run(row1_value)
            self._set_run_font(run, self.size_table_content, bold=False)
            cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            self._set_cell_borders(cell0)
            self._set_cell_borders(cell1)
            self._set_cell_borders(cell2)
            row_idx += 1

        if row2_value:
            # Row 2: Server Tooling/DevOps / Environment
            cell0 = table.cell(row_idx, 0)
            if row_idx == 0:
                cell0.text = ""
                para = cell0.paragraphs[0]
                run = para.add_run(section_title)
                self._set_run_font(run, self.size_table_content, bold=True)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            cell1 = table.cell(row_idx, 1)
            cell1.text = ""
            para = cell1.paragraphs[0]
            run = para.add_run(row2_label)
            self._set_run_font(run, self.size_table_content, bold=False)
            cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            cell2 = table.cell(row_idx, 2)
            cell2.text = ""
            para = cell2.paragraphs[0]
            run = para.add_run(row2_value)
            self._set_run_font(run, self.size_table_content, bold=False)
            cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            self._set_cell_borders(cell0)
            self._set_cell_borders(cell1)
            self._set_cell_borders(cell2)

        # Merge section header cells vertically if multiple rows
        if num_rows > 1:
            self._merge_cells_vertically(table, 0, 0, num_rows - 1)

        doc.add_paragraph()
        return table

    def generate_career_description(
        self,
        data: Dict[str, Any],
        output_path: Path,
        include_personal_info: bool = True
    ) -> int:
        """Generate 경력기술서 document.

        Args:
            data: Template data with user, companies, projects info
            output_path: Path to save the document
            include_personal_info: Whether to include personal info section

        Returns:
            File size in bytes
        """
        doc = Document()

        # Set page margins - exact match to original
        self._set_page_margins(doc)

        # Set default font for document
        style = doc.styles['Normal']
        style.font.name = self.font_body
        style.font.size = Pt(self.size_table_content)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_body)

        # Title: 경력기술서
        self._add_paragraph_with_style(
            doc, "경력기술서",
            self.size_document_title,
            bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=12
        )

        # Name (with English name if available)
        name = data.get("name", "")
        name_en = data.get("name_en", "")
        full_name = f"{name} {name_en}" if name_en else name
        self._add_paragraph_with_style(
            doc, full_name,
            self.size_name,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=6
        )

        # Summary - 8pt as in original
        summary = data.get("summary", "")
        if summary:
            self._add_paragraph_with_style(
                doc, summary,
                self.size_summary,  # 8pt
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=12
            )

        doc.add_paragraph()  # Spacing

        # Personal Info Table (if included)
        if include_personal_info:
            personal_rows = [
                ("성명", data.get("name", "")),
                ("생년월일", data.get("birthdate", "")),
                ("주소", data.get("address", "")),
                ("연락처", data.get("phone", "")),
                ("이메일", data.get("email", "")),
            ]
            self._create_info_table(doc, "인적사항", personal_rows)

        # Education Table
        educations = data.get("educations", [])
        if educations:
            edu_rows = []
            for edu in educations:
                edu_rows.append(("학교명", edu.get("school_name", "")))
                edu_rows.append(("입학일/졸업일", edu.get("period", "")))
                edu_rows.append(("전공", edu.get("major", "")))
                if edu.get("gpa"):
                    edu_rows.append(("학점", edu.get("gpa", "")))
            if edu_rows:
                self._create_info_table(doc, "학력사항", edu_rows)

        # Research/Lab Experience (연구실 경력)
        research = data.get("research", [])
        if research:
            for res in research:
                res_rows = [
                    ("연구실명", res.get("name", "")),
                    ("입학일/퇴학일", res.get("period", "")),
                    ("연구 분야", res.get("research_area", "")),
                ]
                self._create_info_table(doc, "연구실 경력", res_rows, value_bold=(True if res_rows and res_rows[0] else False))

        # Career Table (경력사항)
        companies = data.get("companies", [])
        if companies:
            for company in companies:
                end_date = company.get("end_date", "")
                if not end_date or end_date.lower() == "none":
                    end_date = "현재근무"

                career_rows = [
                    ("회사명", company.get("name", "")),
                    ("입사일/퇴사일", f"{company.get('start_date', '')} ~ {end_date}"),
                    ("소속부서/직급", f"{company.get('department', '')} / {company.get('position', '')}"),
                    ("직무", company.get("description", "")),
                ]
                # Company name should be bold
                self._create_career_table_with_bold_company(doc, "경력사항", career_rows)

        # Freelance Experience (프리랜서 경력)
        freelance = data.get("freelance", [])
        if freelance:
            for fl in freelance:
                end_date = fl.get("end_date", "")
                if not end_date or end_date.lower() == "none":
                    end_date = "종료"

                fl_rows = [
                    ("회사명", fl.get("company_name", "")),
                    ("입사일/퇴사일", f"{fl.get('start_date', '')} ~ {end_date}"),
                    ("소속부서/직급", fl.get("position", "")),
                    ("직무", fl.get("description", "")),
                ]
                self._create_career_table_with_bold_company(doc, "프리랜서 경력", fl_rows)

        # Skills Table (기술스택)
        skills_data = {
            "programming_languages": data.get("programming_languages", ""),
            "frameworks": data.get("frameworks", ""),
            "databases": data.get("databases", ""),
            "cloud": data.get("cloud", ""),
            "tools": data.get("tools", ""),
        }
        self._create_skills_table(doc, skills_data)

        # Projects Table (주요 프로젝트)
        projects = data.get("projects", [])
        if projects:
            self._create_project_table(doc, projects)

        # Certifications
        certifications = data.get("certifications", [])
        if certifications:
            cert_rows = []
            for cert in certifications:
                cert_text = f"{cert.get('name', '')} - {cert.get('issuer', '')} ({cert.get('issue_date', '')})"
                cert_rows.append(("", cert_text))
            if cert_rows:
                self._create_info_table(doc, "자격사항", cert_rows)

        # Awards
        awards = data.get("awards", [])
        if awards:
            award_rows = []
            for award in awards:
                award_text = f"{award.get('name', '')} - {award.get('issuer', '')} ({award.get('award_date', '')})"
                award_rows.append(("", award_text))
            if award_rows:
                self._create_info_table(doc, "수상내역", award_rows)

        # Save document
        os.makedirs(output_path.parent, exist_ok=True)
        doc.save(output_path)
        return os.path.getsize(output_path)

    def generate_resume(
        self,
        data: Dict[str, Any],
        output_path: Path
    ) -> int:
        """Generate 이력서 document (includes salary info and self-introduction).

        Args:
            data: Template data
            output_path: Path to save the document

        Returns:
            File size in bytes
        """
        doc = Document()

        # Set page margins - exact match to original
        self._set_page_margins(doc)

        # Set default font
        style = doc.styles['Normal']
        style.font.name = self.font_body
        style.font.size = Pt(self.size_table_content)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_body)

        # Title: 이력서 (in header)
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run("이력서")
        self._set_run_font(run, self.size_document_title, bold=False)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Name (with English name if available)
        name = data.get("name", "")
        name_en = data.get("name_en", "")
        full_name = f"{name} {name_en}" if name_en else name
        self._add_paragraph_with_style(
            doc, full_name,
            self.size_name,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=6
        )

        # Summary - 8pt as in original
        summary = data.get("summary", "")
        if summary:
            self._add_paragraph_with_style(
                doc, summary,
                self.size_summary,  # 8pt
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=12
            )

        # Salary Info
        salary_info = []
        if data.get("current_salary"):
            salary_info.append(f"현재연봉: {data['current_salary']}")
        if data.get("desired_salary"):
            salary_info.append(f"희망연봉: {data['desired_salary']}")
        if data.get("job_change_reason"):
            salary_info.append(f"이직사유: {data['job_change_reason']}")
        if data.get("available_date"):
            salary_info.append(f"입사 가능일: {data['available_date']}")

        if salary_info:
            for info in salary_info:
                self._add_paragraph_with_style(doc, info, self.size_table_content)
            doc.add_paragraph()

        # Personal Info Table
        personal_rows = [
            ("성명", data.get("name", "")),
            ("생년월일", data.get("birthdate", "")),
            ("주소", data.get("address", "")),
            ("연락처", data.get("phone", "")),
            ("이메일", data.get("email", "")),
        ]
        self._create_info_table(doc, "인적사항", personal_rows)

        # Education Table
        educations = data.get("educations", [])
        if educations:
            edu_rows = []
            for edu in educations:
                edu_rows.append(("학교명", edu.get("school_name", "")))
                edu_rows.append(("입학일/졸업일", edu.get("period", "")))
                edu_rows.append(("전공", edu.get("major", "")))
                if edu.get("gpa"):
                    edu_rows.append(("학점", edu.get("gpa", "")))
            if edu_rows:
                self._create_info_table(doc, "학력사항", edu_rows)

        # Research/Lab Experience (연구실 경력)
        research = data.get("research", [])
        if research:
            for res in research:
                res_rows = [
                    ("연구실명", res.get("name", "")),
                    ("입학일/퇴학일", res.get("period", "")),
                    ("연구 분야", res.get("research_area", "")),
                ]
                self._create_info_table(doc, "연구실 경력", res_rows)

        # Career Table (경력사항) - with bold company name
        companies = data.get("companies", [])
        if companies:
            for company in companies:
                end_date = company.get("end_date", "")
                if not end_date or end_date.lower() == "none":
                    end_date = "현재근무"

                career_rows = [
                    ("회사명", company.get("name", "")),
                    ("입사일/퇴사일", f"{company.get('start_date', '')} ~ {end_date}"),
                    ("소속부서/직급", f"{company.get('department', '')} / {company.get('position', '')}"),
                    ("직무", company.get("description", "")),
                ]
                self._create_career_table_with_bold_company(doc, "경력사항", career_rows)

        # Freelance Experience (프리랜서 경력)
        freelance = data.get("freelance", [])
        if freelance:
            for fl in freelance:
                end_date = fl.get("end_date", "")
                if not end_date or end_date.lower() == "none":
                    end_date = "종료"

                fl_rows = [
                    ("회사명", fl.get("company_name", "")),
                    ("입사일/퇴사일", f"{fl.get('start_date', '')} ~ {end_date}"),
                    ("소속부서/직급", fl.get("position", "")),
                    ("직무", fl.get("description", "")),
                ]
                self._create_career_table_with_bold_company(doc, "프리랜서 경력", fl_rows)

        # Skills Table
        skills_data = {
            "programming_languages": data.get("programming_languages", ""),
            "frameworks": data.get("frameworks", ""),
            "databases": data.get("databases", ""),
            "cloud": data.get("cloud", ""),
            "tools": data.get("tools", ""),
        }
        self._create_skills_table(doc, skills_data)

        # 경력기술서 Section Title - 16pt Bold as in original
        self._add_paragraph_with_style(
            doc, "경력기술서",
            self.size_subsection_title,  # 16pt
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=12
        )

        # Projects Table
        projects = data.get("projects", [])
        if projects:
            self._create_project_table(doc, projects)

        # Certifications
        certifications = data.get("certifications", [])
        if certifications:
            cert_rows = []
            for cert in certifications:
                cert_text = f"{cert.get('name', '')} - {cert.get('issuer', '')} ({cert.get('issue_date', '')})"
                cert_rows.append(("", cert_text))
            if cert_rows:
                self._create_info_table(doc, "자격사항", cert_rows)

        # Awards
        awards = data.get("awards", [])
        if awards:
            award_rows = []
            for award in awards:
                award_text = f"{award.get('name', '')} - {award.get('issuer', '')} ({award.get('award_date', '')})"
                award_rows.append(("", award_text))
            if award_rows:
                self._create_info_table(doc, "수상내역", award_rows)

        # Self Introduction Section (자기소개) - 16pt Bold as in original
        self._add_paragraph_with_style(
            doc, "자기소개",
            self.size_subsection_title,  # 16pt
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=12
        )

        # Self intro table - single cell table matching original
        intro_content = []
        if data.get("motivation"):
            intro_content.append(("지원 동기", data["motivation"]))
        if data.get("competencies"):
            intro_content.append(("업무 수행 역량", data["competencies"]))
        if data.get("personality"):
            intro_content.append(("성격 및 가치관", data["personality"]))

        if intro_content:
            # Create single-cell table for self-intro (width ~15.9cm from original)
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            cell.width = Cm(15.9)

            for idx, (label, content) in enumerate(intro_content):
                if idx == 0:
                    para = cell.paragraphs[0]
                else:
                    para = cell.add_paragraph()
                    para.paragraph_format.space_before = Pt(6)

                run = para.add_run(f"{label}\n")
                self._set_run_font(run, 11, bold=True)

                run = para.add_run(content)
                self._set_run_font(run, self.size_table_content, bold=False)

            self._set_cell_borders(cell)

        # Save document
        os.makedirs(output_path.parent, exist_ok=True)
        doc.save(output_path)
        return os.path.getsize(output_path)
