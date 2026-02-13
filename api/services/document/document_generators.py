"""
Document Generators - Format-specific document generation functions.

Handles DOCX, PDF, Markdown, and HTML generation.
"""
from typing import List, Tuple
from pathlib import Path
import os
import re


def parse_markdown_table(lines: List[str], start_idx: int) -> Tuple[List[List[str]], int]:
    """Parse markdown table starting at given index.

    Returns:
        Tuple of (table_data, end_index) where table_data is list of rows,
        each row is list of cell values.
    """
    table_rows = []
    idx = start_idx

    while idx < len(lines):
        line = lines[idx].strip()
        # Check if line is a table row (starts and ends with |)
        if line.startswith('|') and line.endswith('|'):
            # Parse cells
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            # Skip separator rows (|---|---|)
            if cells and not all(re.match(r'^[-:]+$', cell) for cell in cells):
                table_rows.append(cells)
            idx += 1
        else:
            break

    return table_rows, idx


async def generate_docx(content: str, file_path: Path) -> int:
    """Generate a Word document with proper formatting."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Configure styles - remove colors, set proper font sizes
    # Heading 1: 대제목 (18pt, Bold)
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Malgun Gothic'
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

    # Heading 2: 중제목 (14pt, Bold)
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Malgun Gothic'
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

    # Heading 3: 소제목 (12pt, Bold)
    heading3_style = doc.styles['Heading 3']
    heading3_style.font.name = 'Malgun Gothic'
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True
    heading3_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

    # Normal: 본문 (11pt)
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Malgun Gothic'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

    # List Bullet style
    if 'List Bullet' in doc.styles:
        list_style = doc.styles['List Bullet']
        list_style.font.name = 'Malgun Gothic'
        list_style.font.size = Pt(11)
        list_style.font.color.rgb = RGBColor(0, 0, 0)

    def set_cell_border(cell, border_color="000000", border_size=4):
        """Set cell borders."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(border_size))
            border.set(qn('w:color'), border_color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def style_table_cell(cell, is_header=False, font_size=10):
        """Style a table cell."""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if is_header:
                    run.font.bold = True
        set_cell_border(cell)

    def _add_formatted_paragraph(text: str, style: str = None):
        """Add a paragraph with **bold** markdown parsed into actual bold runs."""
        para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if not part:
                continue
            is_bold = part.startswith('**') and part.endswith('**')
            run = para.add_run(part[2:-2] if is_bold else part)
            if is_bold:
                run.font.bold = True
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
        return para

    # Parse markdown-like content
    lines = content.split('\n')
    idx = 0

    while idx < len(lines):
        line = lines[idx].strip()

        if not line:
            idx += 1
            continue

        # Check if this is the start of a markdown table
        if line.startswith('|') and line.endswith('|'):
            table_data, end_idx = parse_markdown_table(lines, idx)

            if table_data:
                # Determine table dimensions
                num_rows = len(table_data)
                num_cols = max(len(row) for row in table_data) if table_data else 0

                if num_rows > 0 and num_cols > 0:
                    # Create Word table
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Fill in cells
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_value in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_value
                                # First row is usually header
                                is_header = (row_idx == 0)
                                style_table_cell(cell, is_header=is_header)

                    # Add spacing after table
                    doc.add_paragraph()

            idx = end_idx
            continue

        if line.startswith('# '):
            # Heading 1: 대제목 (strip stray ** from heading text)
            para = doc.add_heading(line[2:].replace('**', ''), level=1)
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('## '):
            # Heading 2: 중제목
            para = doc.add_heading(line[3:].replace('**', ''), level=2)
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('### '):
            # Heading 3: 소제목
            para = doc.add_heading(line[4:].replace('**', ''), level=3)
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('#### '):
            # Heading 4: 소소제목 (Bold paragraph)
            para = doc.add_paragraph()
            run = para.add_run(line[5:].replace('**', ''))
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Malgun Gothic'
        elif line.startswith('- ') or line.startswith('• '):
            # Bullet point - parse inline bold
            text = line[2:]
            _add_formatted_paragraph(text, style='List Bullet')
        elif line.startswith('---'):
            # Horizontal line (add empty paragraph)
            doc.add_paragraph()
        else:
            # Regular paragraph - parse inline bold
            _add_formatted_paragraph(line)

        idx += 1

    doc.save(file_path)
    return os.path.getsize(file_path)


async def generate_pdf(content: str, file_path: Path) -> int:
    """Generate a PDF document."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Try to register Korean font
    try:
        # Windows font path
        font_path = "C:/Windows/Fonts/malgun.ttf"
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('MalgunGothic', font_path))
            default_font = 'MalgunGothic'
        else:
            default_font = 'Helvetica'
    except Exception:
        default_font = 'Helvetica'

    doc = SimpleDocTemplate(str(file_path), pagesize=A4)
    styles = getSampleStyleSheet()

    # Create custom styles
    styles.add(ParagraphStyle(
        name='CustomBody',
        fontName=default_font,
        fontSize=11,
        leading=14
    ))
    styles.add(ParagraphStyle(
        name='CustomHeading1',
        fontName=default_font,
        fontSize=18,
        leading=22,
        spaceAfter=12
    ))
    styles.add(ParagraphStyle(
        name='CustomHeading2',
        fontName=default_font,
        fontSize=14,
        leading=18,
        spaceAfter=10
    ))

    elements = []

    lines = content.split('\n')
    for line in lines:
        line = line.strip()

        if line.startswith('# '):
            elements.append(Paragraph(line[2:].replace('**', ''), styles['CustomHeading1']))
        elif line.startswith('## '):
            elements.append(Paragraph(line[3:].replace('**', ''), styles['CustomHeading2']))
        elif line.startswith('### '):
            elements.append(Paragraph(f"<b>{line[4:].replace('**', '')}</b>", styles['CustomBody']))
        elif line.startswith('- '):
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line[2:])
            elements.append(Paragraph(f"• {text}", styles['CustomBody']))
        elif line == '---':
            elements.append(Spacer(1, 0.2 * inch))
        elif line:
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
            elements.append(Paragraph(text, styles['CustomBody']))
        else:
            elements.append(Spacer(1, 0.1 * inch))

    doc.build(elements)
    return os.path.getsize(file_path)


async def generate_markdown(content: str, file_path: Path) -> int:
    """Generate a Markdown document."""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    return os.path.getsize(file_path)


async def generate_html(content: str, file_path: Path) -> int:
    """Generate an HTML document from markdown-like content."""
    import markdown

    # Convert markdown to HTML
    html_body = markdown.markdown(content, extensions=['tables', 'fenced_code', 'toc'])

    # Wrap in a complete HTML document with basic styling
    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume</title>
    <style>
        body {{
            font-family: 'Malgun Gothic', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }}
        h1 {{ font-size: 28px; border-bottom: 2px solid #333; padding-bottom: 10px; }}
        h2 {{ font-size: 20px; color: #0066cc; margin-top: 30px; }}
        h3 {{ font-size: 16px; margin-top: 20px; }}
        ul {{ padding-left: 20px; }}
        li {{ margin-bottom: 5px; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 20px 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
        @media print {{
            body {{ padding: 0; }}
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    return os.path.getsize(file_path)
