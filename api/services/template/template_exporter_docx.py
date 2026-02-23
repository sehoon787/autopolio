"""
DOCX Export Module - Word document generation for template exports

Extracts DOCX-specific logic from TemplateExporter for modularity.
All functions expect a DocxStyler instance for consistent styling.
"""

import os
from pathlib import Path

from api.services.docx.docx_styles import DocxStyler
from api.services.core.mustache_helpers import render_or_clean_mustache
from api.services.export.export_helpers import (
    get_key_tasks_list,
    get_achievements_list,
    group_projects_by_company,
    categorize_skills_by_domain_detailed,
)


def create_docx_from_render_data(
    data,  # RenderDataRequest
    styler: DocxStyler,
    result_dir: Path,
    filename: str,
) -> str:
    """Export to Word document from RenderDataRequest with proper styling

    Uses DocxStyler for consistent formatting:
    - Title (대제목): 24pt, bold, black
    - Heading1 (중제목): 18pt, bold, black
    - Heading2 (소제목): 14pt, bold, black
    - Normal text: 11pt, black

    Args:
        data: RenderDataRequest Pydantic model
        styler: DocxStyler instance for formatting
        result_dir: Output directory path
        filename: Output filename

    Returns:
        Path to saved DOCX file
    """
    from docx import Document

    os.makedirs(result_dir, exist_ok=True)
    file_path = result_dir / filename

    doc = Document()

    # Set up document styles (removes colors, sets fonts)
    styler.setup_document(doc)

    # Header - Title (대제목)
    styler.add_title(doc, data.name or "이력서")

    if data.desired_position:
        styler.add_subtitle(doc, data.desired_position)

    # Contact info
    contact_parts = []
    if data.email:
        contact_parts.append(data.email)
    if data.phone:
        contact_parts.append(data.phone)
    if data.github_url:
        contact_parts.append(data.github_url)

    styler.add_contact_info(doc, contact_parts)
    styler.add_spacing(doc)

    # Summary
    if data.summary:
        styler.add_heading1(doc, "자기소개")
        styler.add_paragraph(doc, data.summary)

    # Experience
    if data.experiences:
        styler.add_heading1(doc, "경력사항")
        for exp in data.experiences:
            # Company name as heading2 (소제목)
            period = f" ({exp.start_date or ''} ~ {exp.end_date or '현재'})"
            styler.add_heading2(doc, f"{exp.company_name}{period}")

            # Build context for Mustache rendering
            exp_context = {
                "company_name": exp.company_name,
                "position": exp.position,
                "start_date": exp.start_date,
                "end_date": exp.end_date,
                "description": exp.description,
                "achievements": exp.achievements,
            }

            if exp.position:
                styler.add_bullet(doc, exp.position)
            if exp.description:
                # Clean any Mustache syntax from description
                clean_desc = render_or_clean_mustache(exp.description, exp_context)
                if clean_desc:
                    styler.add_paragraph(doc, clean_desc)
            if exp.achievements:
                for ach in exp.achievements:
                    # Clean any Mustache syntax from achievements
                    clean_ach = render_or_clean_mustache(ach, exp_context)
                    if clean_ach:
                        styler.add_bullet(doc, clean_ach)

    # Projects
    if data.projects:
        styler.add_heading1(doc, "프로젝트")
        for proj in data.projects:
            # Project name as heading2 (소제목)
            period = f" ({proj.start_date or ''} ~ {proj.end_date or '현재'})"
            styler.add_heading2(doc, f"{proj.name}{period}")

            # Build context for Mustache rendering
            proj_context = {
                "name": proj.name,
                "company_name": proj.company_name,
                "start_date": proj.start_date,
                "end_date": proj.end_date,
                "description": proj.description,
                "role": proj.role,
                "technologies": proj.technologies,
                "achievements": proj.achievements,
            }

            if proj.company_name:
                styler.add_paragraph(doc, f"소속: {proj.company_name}")
            if proj.description:
                # Clean any Mustache syntax from description
                clean_desc = render_or_clean_mustache(proj.description, proj_context)
                if clean_desc:
                    styler.add_paragraph(doc, clean_desc)
            if proj.role:
                # Clean any Mustache syntax from role
                clean_role = render_or_clean_mustache(proj.role, proj_context)
                if clean_role:
                    styler.add_paragraph(doc, f"역할: {clean_role}")
            if proj.technologies:
                styler.add_paragraph(doc, f"기술: {', '.join(proj.technologies)}")
            # Use achievements_summary_list (from detailed_achievements) as default, fall back to achievements
            achievements_list = get_achievements_list(
                {
                    "achievements_summary_list": proj.achievements_summary_list,
                    "achievements_detailed_list": proj.achievements_detailed_list,
                    "achievements": proj.achievements,
                }
            )
            if achievements_list:
                for ach in achievements_list:
                    # Clean any Mustache syntax from achievements
                    clean_ach = render_or_clean_mustache(ach, proj_context)
                    if clean_ach:
                        styler.add_bullet(doc, clean_ach)

    # Skills
    if data.skills:
        styler.add_heading1(doc, "기술 스택")
        skills = data.skills
        if skills.languages:
            styler.add_key_value(doc, "Languages", ", ".join(skills.languages))
        if skills.frameworks:
            styler.add_key_value(doc, "Frameworks", ", ".join(skills.frameworks))
        if skills.databases:
            styler.add_key_value(doc, "Databases", ", ".join(skills.databases))
        if skills.tools:
            styler.add_key_value(doc, "Tools", ", ".join(skills.tools))

    # Education
    if data.educations:
        styler.add_heading1(doc, "학력")
        for edu in data.educations:
            period = f" ({edu.start_date or ''} ~ {edu.end_date or ''})"
            major_str = f" - {edu.major}" if edu.major else ""
            styler.add_heading2(doc, f"{edu.school_name}{major_str}{period}")

    # Certifications
    if data.certifications:
        styler.add_heading1(doc, "자격증")
        for cert in data.certifications:
            styler.add_bullet(
                doc, f"{cert.name} ({cert.issuer or ''}) - {cert.date or ''}"
            )

    doc.save(file_path)
    return str(file_path)


def create_unified_docx(
    data: dict, styler: DocxStyler, result_dir: Path, filename: str
) -> str:
    """Unified Word format for all platforms with proper styling

    Structure:
    - 경력: Company-based with domain-categorized skills
    - 경력기술서: Detailed project information

    Styling:
    - Title (대제목): 24pt, bold, black
    - Heading1 (중제목): 18pt, bold, black
    - Heading2 (소제목): 14pt, bold, black
    - Normal text: 11pt, black

    Args:
        data: User data dictionary
        styler: DocxStyler instance for formatting
        result_dir: Output directory path
        filename: Output filename

    Returns:
        Path to saved DOCX file
    """
    from docx import Document

    os.makedirs(result_dir, exist_ok=True)
    file_path = result_dir / filename

    doc = Document()

    # Set up document styles (removes colors, sets fonts)
    styler.setup_document(doc)

    # Header - Title (대제목)
    styler.add_title(doc, data.get("name", "이력서"))

    # Contact info
    _add_contact_info_docx(doc, data, styler)

    experiences = data.get("experiences", [])
    projects = data.get("projects", [])

    # Calculate total experience
    total_years = data.get("total_experience_years", 0)

    # === Part 1: 경력 ===
    if experiences or projects:
        heading_text = f"경력 (총 {total_years}년)" if total_years else "경력"
        styler.add_heading1(doc, heading_text)

        company_projects = group_projects_by_company(projects)

        for exp in experiences:
            company_name = exp.get("company_name", "")
            position = exp.get("position", "")
            start_date = exp.get("start_date", "")
            end_date = exp.get("end_date", "")
            is_current = exp.get("is_current", not end_date)
            duration = exp.get("duration", "")
            description = exp.get("description", "")

            # Company header (소제목)
            styler.add_heading2(doc, company_name)

            if position:
                styler.add_paragraph(doc, position)

            period_str = f"{start_date} ~ {end_date if end_date else ''}"
            if is_current:
                period_str += " 재직중" if not end_date else ""
            if duration:
                period_str += f" ({duration})"
            styler.add_paragraph(doc, period_str)

            if description:
                styler.add_bullet(doc, f"주요 직무: {description}")

            # Domain skills
            company_projs = company_projects.get(company_name, [])
            skills_by_domain = categorize_skills_by_domain_detailed(company_projs)

            domain_idx = 1
            for domain_name, domain_data in skills_by_domain.items():
                if not domain_data.get("technologies") and not domain_data.get(
                    "implementations"
                ):
                    continue

                styler.add_bold_text(doc, f"{domain_idx}. {domain_name}")

                if domain_data.get("technologies"):
                    styler.add_paragraph(doc, "주요 사용 기술")
                    for tech_line in domain_data["technologies"]:
                        styler.add_bullet(doc, tech_line)

                if domain_data.get("implementations"):
                    styler.add_paragraph(doc, "구현 내용")
                    for impl in domain_data["implementations"]:
                        styler.add_bullet(doc, impl)

                if domain_data.get("databases"):
                    styler.add_paragraph(doc, "DB")
                    for db_line in domain_data["databases"]:
                        styler.add_bullet(doc, db_line)

                domain_idx += 1

            styler.add_spacing(doc)

        # Side projects
        side_projects = (
            company_projects.get(None, [])
            + company_projects.get("", [])
            + company_projects.get("사이드 프로젝트", [])
        )
        if side_projects:
            styler.add_heading2(doc, "사이드 프로젝트")

            skills_by_domain = categorize_skills_by_domain_detailed(side_projects)

            domain_idx = 1
            for domain_name, domain_data in skills_by_domain.items():
                if not domain_data.get("implementations"):
                    continue

                styler.add_bold_text(doc, f"{domain_idx}. {domain_name}")

                for impl in domain_data["implementations"]:
                    styler.add_bullet(doc, impl)

                domain_idx += 1

    # === Part 2: 경력기술서 ===
    if projects:
        styler.add_heading1(doc, "경력기술서")

        for idx, proj in enumerate(projects, 1):
            _add_project_detail_unified_docx(doc, proj, idx, styler)

    doc.save(file_path)
    return str(file_path)


def _add_contact_info_docx(doc, data: dict, styler: DocxStyler) -> None:
    """Add contact info to Word document using styler"""
    contact_parts = []
    if data.get("email"):
        contact_parts.append(data["email"])
    if data.get("phone"):
        contact_parts.append(data["phone"])
    if data.get("github_url"):
        contact_parts.append(data["github_url"])

    styler.add_contact_info(doc, contact_parts)
    styler.add_spacing(doc)


def _add_project_detail_unified_docx(
    doc, proj: dict, idx: int, styler: DocxStyler
) -> None:
    """Add project in unified format to Word document using styler

    Note: Uses render_or_clean_mustache to handle any Mustache template
    syntax that may be in the text fields (e.g., {{key_tasks}}, {{#achievements}})
    """
    # Project header (소소제목)
    styler.add_section_title(doc, f"[프로젝트 {idx}]")

    styler.add_key_value(doc, "프로젝트명", proj.get("name", ""))

    company = proj.get("company_name") or proj.get("company")
    if company:
        styler.add_key_value(doc, "연계/소속회사", company)
    else:
        styler.add_key_value(doc, "연계/소속회사", "사이드 프로젝트")

    end_date = proj.get("end_date") or "진행중"
    styler.add_key_value(doc, "기간", f"{proj.get('start_date', '')} ~ {end_date}")

    role = proj.get("role")
    if role:
        # Clean any Mustache syntax from role
        clean_role = render_or_clean_mustache(role, proj)
        if clean_role:
            styler.add_key_value(doc, "주요 업무", clean_role)

    # 주요 구현 기능 (최대 8개)
    key_tasks = get_key_tasks_list(proj)
    if key_tasks:
        styler.add_paragraph(doc, "주요 구현 기능:")
        for task in key_tasks[:8]:
            # Clean any Mustache syntax from each task
            clean_task = render_or_clean_mustache(task, proj)
            if clean_task:
                styler.add_bullet(doc, clean_task)

    # 기술 스택
    tech = proj.get("technologies")
    if tech:
        tech_str = tech if isinstance(tech, str) else ", ".join(tech)
        styler.add_key_value(doc, "기술 스택", tech_str)

    # 개발 인원
    team_size = proj.get("team_size")
    if team_size:
        styler.add_key_value(doc, "개발 인원", str(team_size))

    # 상세 내용 - clean any Mustache template syntax
    description = proj.get("description", "")
    if description:
        # Render or clean Mustache syntax from description
        clean_description = render_or_clean_mustache(description, proj)
        if clean_description:
            styler.add_key_value(doc, "상세 내용", clean_description)

    # 성과 (grouped by category)
    achievements_grouped = proj.get("achievements_grouped", [])
    if achievements_grouped:
        styler.add_paragraph(doc, "성과:")
        for group in achievements_grouped:
            category = group.get("category", "")
            items = group.get("items", [])
            if category:
                styler.add_bold_text(doc, f"[{category}]")
            for item in items:
                title = item.get("title", "") if isinstance(item, dict) else str(item)
                if title:
                    styler.add_bullet(doc, title)
    else:
        # Fallback to flat list
        achievements = get_achievements_list(proj)
        if achievements:
            styler.add_paragraph(doc, "성과:")
            for ach in achievements:
                clean_ach = render_or_clean_mustache(ach, proj)
                if clean_ach:
                    styler.add_bullet(doc, clean_ach)

    styler.add_spacing(doc)
