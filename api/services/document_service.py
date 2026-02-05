"""
Document Service - Handles document generation (DOCX, PDF, Markdown).
"""
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import os
import re
from datetime import datetime

from api.config import get_settings

settings = get_settings()


class DocumentService:
    """Service for document generation and template parsing."""

    def __init__(self):
        self.result_dir = settings.result_dir

    async def parse_template(self, file_path: str) -> Tuple[str, Dict[str, str]]:
        """Parse a template file and extract field mappings."""
        ext = Path(file_path).suffix.lower()

        if ext in [".docx", ".doc"]:
            return await self._parse_docx_template(file_path)
        elif ext == ".pdf":
            return await self._parse_pdf_template(file_path)
        else:
            raise ValueError(f"Unsupported template format: {ext}")

    async def _parse_docx_template(self, file_path: str) -> Tuple[str, Dict[str, str]]:
        """Parse a Word document template."""
        from docx import Document

        doc = Document(file_path)
        content_parts = []
        field_mappings = {}

        # Pattern for placeholders like {{field_name}} or {field_name}
        placeholder_pattern = r'\{\{?\s*([a-zA-Z_][a-zA-Z0-9_\.]*)\s*\}?\}'

        for para in doc.paragraphs:
            text = para.text
            content_parts.append(text)

            # Find all placeholders
            matches = re.findall(placeholder_pattern, text)
            for match in matches:
                field_mappings[f"{{{{{match}}}}}"] = match

        full_content = "\n".join(content_parts)
        return full_content, field_mappings

    async def _parse_pdf_template(self, file_path: str) -> Tuple[str, Dict[str, str]]:
        """Parse a PDF template (extract text and find placeholders)."""
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        content_parts = []
        field_mappings = {}

        placeholder_pattern = r'\{\{?\s*([a-zA-Z_][a-zA-Z0-9_\.]*)\s*\}?\}'

        for page in reader.pages:
            text = page.extract_text()
            content_parts.append(text)

            matches = re.findall(placeholder_pattern, text)
            for match in matches:
                field_mappings[f"{{{{{match}}}}}"] = match

        full_content = "\n".join(content_parts)
        return full_content, field_mappings

    async def generate_document(
        self,
        template_content: str,
        data: Dict[str, Any],
        output_format: str,
        document_name: str,
        template_platform: str = None
    ) -> Tuple[str, int]:
        """Generate a document from template and data.

        Args:
            template_content: Template content (markdown/mustache format)
            data: Data to fill into template
            output_format: Output format (docx, pdf, md, html)
            document_name: Name for the output file
            template_platform: Platform type for special handling (career_description, resume, etc.)
        """
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r'[^\w\-_]', '_', document_name)
        filename = f"{safe_name}_{timestamp}.{output_format}"
        file_path = self.result_dir / filename

        os.makedirs(self.result_dir, exist_ok=True)

        # Use specialized generator for Korean resume templates
        if output_format == "docx" and template_platform in (
            "career_description",
            "career_description_no_personal",
            "resume"
        ):
            from api.services.docx_generator import DocxGenerator
            generator = DocxGenerator()

            if template_platform == "career_description":
                file_size = generator.generate_career_description(
                    data, file_path, include_personal_info=True
                )
            elif template_platform == "career_description_no_personal":
                file_size = generator.generate_career_description(
                    data, file_path, include_personal_info=False
                )
            elif template_platform == "resume":
                file_size = generator.generate_resume(data, file_path)
            else:
                file_size = generator.generate_career_description(data, file_path)

            return str(file_path), file_size

        # For other templates, use markdown-based rendering
        rendered_content = self._render_template(template_content, data)

        if output_format == "docx":
            file_size = await self._generate_docx(rendered_content, file_path)
        elif output_format == "pdf":
            file_size = await self._generate_pdf(rendered_content, file_path)
        elif output_format == "md":
            file_size = await self._generate_markdown(rendered_content, file_path)
        elif output_format == "html":
            file_size = await self._generate_html(rendered_content, file_path)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")

        return str(file_path), file_size

    def _render_template(self, template: str, data: Dict[str, Any]) -> str:
        """Render template with data (simple mustache-like syntax)."""
        result = template

        # Handle simple variables {{variable}}
        for key, value in self._flatten_dict(data).items():
            placeholder = f"{{{{{key}}}}}"
            if isinstance(value, list):
                # If it's a list of simple strings, render as bullet points
                if all(isinstance(v, str) for v in value):
                    value = "\n".join(f"- {v}" for v in value) if value else ""
                else:
                    # For lists of dicts, just convert to string (sections handle these)
                    value = ", ".join(str(v) for v in value)
            elif value is None:
                value = ""
            result = result.replace(placeholder, str(value))

        # Handle sections {{#section}}...{{/section}}
        result = self._render_sections(result, data)

        return result

    def _flatten_dict(self, d: Dict, parent_key: str = '', sep: str = '.') -> Dict:
        """Flatten nested dictionary."""
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_dict(v, new_key, sep).items())
            else:
                items.append((new_key, v))
        return dict(items)

    def _render_sections(self, template: str, data: Dict[str, Any]) -> str:
        """Render section loops in template with recursive support for nested sections."""
        # Pattern for sections: {{#section}}...{{/section}}
        # Using non-greedy matching for innermost sections first
        section_pattern = r'\{\{#(\w+)\}\}(.*?)\{\{/\1\}\}'

        def render_item(item_data: Dict[str, Any], content: str) -> str:
            """Recursively render a single item with its nested sections."""
            result = content

            # First, handle boolean conditional sections at the current item level
            # e.g., {{#has_description}}...{{/has_description}}
            for key, value in item_data.items():
                if isinstance(value, bool):
                    bool_pattern = rf'\{{\{{#{key}\}}\}}(.*?)\{{\{{/{key}\}}\}}'
                    if value:
                        # True - keep the content inside
                        result = re.sub(bool_pattern, r'\1', result, flags=re.DOTALL)
                    else:
                        # False - remove the entire section
                        result = re.sub(bool_pattern, '', result, flags=re.DOTALL)

            # Then, recursively process any nested sections within this item
            for key, value in item_data.items():
                if isinstance(value, list) and len(value) > 0 and isinstance(value[0], dict):
                    # This is a nested section (like achievements inside projects)
                    nested_pattern = rf'\{{\{{#{key}\}}\}}(.*?)\{{\{{/{key}\}}\}}'
                    nested_match = re.search(nested_pattern, result, flags=re.DOTALL)
                    if nested_match:
                        nested_content = nested_match.group(1)
                        nested_rendered = []
                        for nested_item in value:
                            nested_item_content = nested_content
                            # First handle boolean conditional sections (e.g., {{#has_description}}...{{/has_description}})
                            for nkey, nvalue in nested_item.items():
                                if isinstance(nvalue, bool):
                                    bool_pattern = rf'\{{\{{#{nkey}\}}\}}(.*?)\{{\{{/{nkey}\}}\}}'
                                    if nvalue:
                                        # True - keep the content inside
                                        nested_item_content = re.sub(bool_pattern, r'\1', nested_item_content, flags=re.DOTALL)
                                    else:
                                        # False - remove the entire section
                                        nested_item_content = re.sub(bool_pattern, '', nested_item_content, flags=re.DOTALL)
                            # Then replace simple placeholders
                            for nkey, nvalue in nested_item.items():
                                placeholder = f"{{{{{nkey}}}}}"
                                nested_item_content = nested_item_content.replace(
                                    placeholder, str(nvalue if nvalue is not None and not isinstance(nvalue, bool) else '')
                                )
                            nested_rendered.append(nested_item_content.strip())
                        result = re.sub(nested_pattern, '\n'.join(nested_rendered), result, flags=re.DOTALL)

            # Then replace simple field placeholders
            for key, value in item_data.items():
                placeholder = f"{{{{{key}}}}}"
                if isinstance(value, list):
                    # Convert list to bullet points for simple string lists
                    if value and all(isinstance(v, str) for v in value):
                        list_str = "\n".join(f"- {v}" for v in value)
                        result = result.replace(placeholder, list_str)
                    elif not value:
                        # Empty list - remove placeholder
                        result = result.replace(placeholder, "")
                    # Lists of dicts are handled by nested section logic above
                elif not isinstance(value, dict):
                    result = result.replace(placeholder, str(value or ''))

            return result

        def replace_section(match):
            section_name = match.group(1)
            section_content = match.group(2)

            items = data.get(section_name, [])
            if not isinstance(items, list):
                items = [items] if items else []

            rendered_items = []
            for item in items:
                if isinstance(item, dict):
                    item_content = render_item(item, section_content)
                else:
                    item_content = section_content.replace("{{.}}", str(item))
                rendered_items.append(item_content.strip())

            return "\n".join(rendered_items)

        # Apply section replacement - may need multiple passes for complex nesting
        result = template
        prev_result = None
        max_iterations = 10  # Prevent infinite loops
        iterations = 0

        while result != prev_result and iterations < max_iterations:
            prev_result = result
            result = re.sub(section_pattern, replace_section, result, flags=re.DOTALL)
            iterations += 1

        return result

    def _parse_markdown_table(self, lines: List[str], start_idx: int) -> Tuple[List[List[str]], int]:
        """Parse markdown table starting at given index.

        Returns:
            Tuple of (table_data, end_index) where table_data is list of rows,
            each row is list of cell values.
        """
        table_rows = []
        idx = start_idx

        while idx < len(lines):
            line = lines[idx].strip()
            # Check if line is a table row (starts and ends with |)
            if line.startswith('|') and line.endswith('|'):
                # Parse cells
                cells = [cell.strip() for cell in line[1:-1].split('|')]
                # Skip separator rows (|---|---|)
                if cells and not all(re.match(r'^[-:]+$', cell) for cell in cells):
                    table_rows.append(cells)
                idx += 1
            else:
                break

        return table_rows, idx

    async def _generate_docx(self, content: str, file_path: Path) -> int:
        """Generate a Word document with proper formatting."""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Configure styles - remove colors, set proper font sizes
        # Heading 1: 대제목 (18pt, Bold)
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = 'Malgun Gothic'
        heading1_style.font.size = Pt(18)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Heading 2: 중제목 (14pt, Bold)
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Malgun Gothic'
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True
        heading2_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Heading 3: 소제목 (12pt, Bold)
        heading3_style = doc.styles['Heading 3']
        heading3_style.font.name = 'Malgun Gothic'
        heading3_style.font.size = Pt(12)
        heading3_style.font.bold = True
        heading3_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Normal: 본문 (11pt)
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Malgun Gothic'
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # List Bullet style
        if 'List Bullet' in doc.styles:
            list_style = doc.styles['List Bullet']
            list_style.font.name = 'Malgun Gothic'
            list_style.font.size = Pt(11)
            list_style.font.color.rgb = RGBColor(0, 0, 0)

        def set_cell_border(cell, border_color="000000", border_size=4):
            """Set cell borders."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(border_size))
                border.set(qn('w:color'), border_color)
                tcBorders.append(border)
            tcPr.append(tcBorders)

        def style_table_cell(cell, is_header=False, font_size=10):
            """Style a table cell."""
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Malgun Gothic'
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if is_header:
                        run.font.bold = True
            set_cell_border(cell)

        # Parse markdown-like content
        lines = content.split('\n')
        idx = 0

        while idx < len(lines):
            line = lines[idx].strip()

            if not line:
                idx += 1
                continue

            # Check if this is the start of a markdown table
            if line.startswith('|') and line.endswith('|'):
                table_data, end_idx = self._parse_markdown_table(lines, idx)

                if table_data:
                    # Determine table dimensions
                    num_rows = len(table_data)
                    num_cols = max(len(row) for row in table_data) if table_data else 0

                    if num_rows > 0 and num_cols > 0:
                        # Create Word table
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                        # Fill in cells
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_value in enumerate(row_data):
                                if col_idx < num_cols:
                                    cell = table.cell(row_idx, col_idx)
                                    cell.text = cell_value
                                    # First row is usually header
                                    is_header = (row_idx == 0)
                                    style_table_cell(cell, is_header=is_header)

                        # Add spacing after table
                        doc.add_paragraph()

                idx = end_idx
                continue

            if line.startswith('# '):
                # Heading 1: 대제목
                para = doc.add_heading(line[2:], level=1)
                # Ensure color is black (override theme)
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            elif line.startswith('## '):
                # Heading 2: 중제목
                para = doc.add_heading(line[3:], level=2)
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            elif line.startswith('### '):
                # Heading 3: 소제목
                para = doc.add_heading(line[4:], level=3)
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            elif line.startswith('#### '):
                # Heading 4: 소소제목 (Bold paragraph)
                para = doc.add_paragraph()
                run = para.add_run(line[5:])
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Malgun Gothic'
            elif line.startswith('- ') or line.startswith('• '):
                # Bullet point
                text = line[2:] if line.startswith('- ') else line[2:]
                para = doc.add_paragraph(text, style='List Bullet')
            elif line.startswith('---'):
                # Horizontal line (add empty paragraph)
                doc.add_paragraph()
            else:
                # Regular paragraph
                para = doc.add_paragraph(line)

            idx += 1

        doc.save(file_path)
        return os.path.getsize(file_path)

    async def _generate_pdf(self, content: str, file_path: Path) -> int:
        """Generate a PDF document."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # Try to register Korean font
        try:
            # Windows font path
            font_path = "C:/Windows/Fonts/malgun.ttf"
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('MalgunGothic', font_path))
                default_font = 'MalgunGothic'
            else:
                default_font = 'Helvetica'
        except Exception:
            default_font = 'Helvetica'

        doc = SimpleDocTemplate(str(file_path), pagesize=A4)
        styles = getSampleStyleSheet()

        # Create custom styles
        styles.add(ParagraphStyle(
            name='CustomBody',
            fontName=default_font,
            fontSize=11,
            leading=14
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading1',
            fontName=default_font,
            fontSize=18,
            leading=22,
            spaceAfter=12
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading2',
            fontName=default_font,
            fontSize=14,
            leading=18,
            spaceAfter=10
        ))

        elements = []

        lines = content.split('\n')
        for line in lines:
            line = line.strip()

            if line.startswith('# '):
                elements.append(Paragraph(line[2:], styles['CustomHeading1']))
            elif line.startswith('## '):
                elements.append(Paragraph(line[3:], styles['CustomHeading2']))
            elif line.startswith('### '):
                elements.append(Paragraph(f"<b>{line[4:]}</b>", styles['CustomBody']))
            elif line.startswith('- '):
                elements.append(Paragraph(f"• {line[2:]}", styles['CustomBody']))
            elif line == '---':
                elements.append(Spacer(1, 0.2 * inch))
            elif line:
                elements.append(Paragraph(line, styles['CustomBody']))
            else:
                elements.append(Spacer(1, 0.1 * inch))

        doc.build(elements)
        return os.path.getsize(file_path)

    async def _generate_markdown(self, content: str, file_path: Path) -> int:
        """Generate a Markdown document."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return os.path.getsize(file_path)

    async def _generate_html(self, content: str, file_path: Path) -> int:
        """Generate an HTML document from markdown-like content."""
        import markdown

        # Convert markdown to HTML
        html_body = markdown.markdown(content, extensions=['tables', 'fenced_code', 'toc'])

        # Wrap in a complete HTML document with basic styling
        html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume</title>
    <style>
        body {{
            font-family: 'Malgun Gothic', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }}
        h1 {{ font-size: 28px; border-bottom: 2px solid #333; padding-bottom: 10px; }}
        h2 {{ font-size: 20px; color: #0066cc; margin-top: 30px; }}
        h3 {{ font-size: 16px; margin-top: 20px; }}
        ul {{ padding-left: 20px; }}
        li {{ margin-bottom: 5px; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 20px 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
        @media print {{
            body {{ padding: 0; }}
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return os.path.getsize(file_path)

    def prepare_template_data(
        self,
        user_data: Dict[str, Any],
        companies: List[Dict[str, Any]],
        projects: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Prepare data structure for template rendering."""
        # Extract categorized skills from all projects
        skill_categories = self._categorize_skills(projects)

        return {
            # Basic user info
            "name": user_data.get("name", ""),
            "email": user_data.get("email", ""),
            "github_username": user_data.get("github_username", ""),
            "summary": user_data.get("summary", ""),

            # Personal information (from user profile)
            "birthdate": user_data.get("birthdate", ""),
            "address": user_data.get("address", ""),
            "phone": user_data.get("phone", ""),
            "military_status": user_data.get("military_status", ""),
            "military_type": user_data.get("military_type", ""),

            # Education
            "university": user_data.get("university", ""),
            "major": user_data.get("major", ""),
            "graduation_date": user_data.get("graduation_date", ""),
            "gpa": user_data.get("gpa", ""),

            # Categorized skills
            "programming_languages": skill_categories.get("programming_languages", ""),
            "frameworks": skill_categories.get("frameworks", ""),
            "databases": skill_categories.get("databases", ""),
            "tools": skill_categories.get("tools", ""),
            "cloud": skill_categories.get("cloud", ""),
            "devops": skill_categories.get("cloud", ""),  # Alias for templates using devops
            "tooling": skill_categories.get("tools", ""),  # Alias for templates using tooling

            # Legacy skills (all combined)
            "skills": self._extract_all_skills(projects),

            # Salary and job change info (from user profile)
            "current_salary": user_data.get("current_salary", ""),
            "desired_salary": user_data.get("desired_salary", ""),
            "job_change_reason": user_data.get("job_change_reason", ""),
            "available_date": user_data.get("available_date", ""),
            "motivation": user_data.get("motivation", ""),
            "competencies": user_data.get("competencies", ""),
            "personality": user_data.get("personality", ""),

            # Companies/Career
            "companies": [
                {
                    "name": c.get("name", ""),
                    "position": c.get("position", ""),
                    "department": c.get("department", ""),
                    "start_date": str(c.get("start_date", "")),
                    "end_date": str(c.get("end_date", "")) if c.get("end_date") else "현재",
                    "description": c.get("description", ""),
                    "location": c.get("location", "")
                }
                for c in companies
            ],

            # Build company lookup for project enrichment
            **self._build_project_company_data(projects, companies),

            # Total career experience
            "total_experience": self._calculate_total_experience(companies)
        }

    def _categorize_skills(self, projects: List[Dict[str, Any]]) -> Dict[str, str]:
        """Categorize skills from projects into predefined categories."""
        # Skill category mappings
        categories = {
            "programming_languages": {
                "Python", "JavaScript", "TypeScript", "Java", "Kotlin", "Go", "Rust",
                "C", "C++", "C#", "PHP", "Ruby", "Swift", "Dart", "Scala", "R",
                "MATLAB", "Perl", "Shell", "Bash", "PowerShell", "SQL", "HTML", "CSS"
            },
            "frameworks": {
                "React", "Vue", "Angular", "Next.js", "Nuxt.js", "Svelte", "Django",
                "Flask", "FastAPI", "Spring", "Spring Boot", "Express", "NestJS",
                "Rails", "Laravel", "ASP.NET", "Flutter", "React Native", "Electron",
                "Tailwind CSS", "Bootstrap", "Material-UI", "MUI", "Ant Design",
                "Redux", "MobX", "Zustand", "TanStack Query", "jQuery"
            },
            "databases": {
                "PostgreSQL", "MySQL", "MariaDB", "MongoDB", "Redis", "SQLite",
                "Oracle", "SQL Server", "DynamoDB", "Cassandra", "Elasticsearch",
                "InfluxDB", "Neo4j", "Firebase", "Supabase", "CockroachDB"
            },
            "tools": {
                "Git", "GitHub", "GitLab", "Bitbucket", "Docker", "Kubernetes",
                "Jenkins", "CircleCI", "GitHub Actions", "Terraform", "Ansible",
                "Webpack", "Vite", "Babel", "ESLint", "Prettier", "Jest", "Pytest",
                "Playwright", "Cypress", "Selenium", "Postman", "Swagger", "Grafana",
                "Prometheus", "Nginx", "Apache", "VS Code", "IntelliJ", "Figma"
            },
            "cloud": {
                "AWS", "GCP", "Google Cloud", "Azure", "Heroku", "Vercel", "Netlify",
                "DigitalOcean", "Cloudflare", "Firebase", "Supabase", "Render"
            }
        }

        result = {cat: set() for cat in categories}

        for project in projects:
            techs = project.get("technologies", [])
            if isinstance(techs, str):
                techs = [t.strip() for t in techs.split(",")]

            for tech in techs:
                tech_normalized = tech.strip()
                for category, keywords in categories.items():
                    if tech_normalized in keywords:
                        result[category].add(tech_normalized)
                        break

        return {cat: ", ".join(sorted(skills)) for cat, skills in result.items()}

    def _build_project_company_data(
        self,
        projects: List[Dict[str, Any]],
        companies: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Build project data enriched with company info."""
        # Build company lookup by ID
        company_lookup = {c.get("id"): c for c in companies if c.get("id")}

        enriched_projects = []
        for p in projects:
            # Get company info for project
            company_id = p.get("company_id")
            company = company_lookup.get(company_id, {}) if company_id else {}

            # Use achievements_summary_list as default (title only, no description)
            # achievements_detailed_list is kept for templates that need description
            raw_achievements = p.get("achievements", [])
            achievements_summary_list = p.get("achievements_summary_list", [])
            achievements_detailed_list = p.get("achievements_detailed_list", [])

            # If achievements lists not provided, build from raw achievements
            if not achievements_summary_list and not achievements_detailed_list and raw_achievements:
                achievements_summary_list = []
                achievements_detailed_list = []
                for a in raw_achievements:
                    # Check if already in correct format (has 'category' and 'title')
                    if "category" in a and "title" in a:
                        # Summary: title only, no description
                        achievements_summary_list.append({
                            "category": a.get("category", "성과"),
                            "title": a.get("title", ""),
                        })
                        # Detailed: with description
                        achievements_detailed_list.append({
                            "category": a.get("category", "성과"),
                            "title": a.get("title", ""),
                            "description": a.get("description", ""),
                            "has_description": bool(a.get("description"))
                        })
                    else:
                        # Legacy format: metric_name/metric_value
                        display_category = a.get("metric_name", "성과")
                        title = a.get("metric_value", "")
                        description = a.get("description", "")
                        # Summary: title only
                        achievements_summary_list.append({
                            "category": display_category,
                            "title": title,
                        })
                        # Detailed: with description
                        achievements_detailed_list.append({
                            "category": display_category,
                            "title": title,
                            "description": description,
                            "has_description": bool(description)
                        })

            enriched_projects.append({
                "name": p.get("name", ""),
                "short_description": p.get("short_description", ""),
                "description": p.get("description", "") or p.get("ai_summary", ""),
                "start_date": str(p.get("start_date", "")),
                "end_date": str(p.get("end_date", "")) if p.get("end_date") else "진행중",
                "role": p.get("role", ""),
                "team_size": p.get("team_size", ""),
                "contribution_percent": p.get("contribution_percent", ""),
                "technologies": ", ".join(p.get("technologies", [])) if isinstance(p.get("technologies"), list) else p.get("technologies", ""),
                "key_tasks": p.get("key_tasks", []),
                "achievements": raw_achievements,  # Keep original for backward compatibility
                "achievements_summary_list": achievements_summary_list,  # Default: title only, no description
                "achievements_detailed_list": achievements_detailed_list,  # With description
                "links": p.get("links", {}),
                # Company info for project
                "company_name": company.get("name", ""),
                "department": company.get("department", ""),
                "position": company.get("position", ""),
            })

        return {"projects": enriched_projects}

    def _calculate_total_experience(self, companies: List[Dict[str, Any]]) -> str:
        """Calculate total career experience from companies."""
        if not companies:
            return ""

        total_months = 0
        for company in companies:
            start = company.get("start_date")
            end = company.get("end_date")

            if start:
                try:
                    if isinstance(start, str):
                        start_date = datetime.strptime(start[:10], "%Y-%m-%d")
                    else:
                        start_date = start

                    if end:
                        if isinstance(end, str):
                            end_date = datetime.strptime(end[:10], "%Y-%m-%d")
                        else:
                            end_date = end
                    else:
                        end_date = datetime.now()

                    months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
                    total_months += max(0, months)
                except (ValueError, TypeError):
                    continue

        years = total_months // 12
        months = total_months % 12

        if years > 0 and months > 0:
            return f"{years}년 {months}개월"
        elif years > 0:
            return f"{years}년"
        elif months > 0:
            return f"{months}개월"
        return ""

    def _extract_all_skills(self, projects: List[Dict[str, Any]]) -> str:
        """Extract and deduplicate all skills from projects."""
        all_techs = set()
        for project in projects:
            techs = project.get("technologies", [])
            if isinstance(techs, str):
                techs = [t.strip() for t in techs.split(",")]
            all_techs.update(techs)
        return ", ".join(sorted(all_techs))
